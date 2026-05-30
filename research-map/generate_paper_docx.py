"""
Generate Paper-Tofu-Final-IOP-EES.docx -- IOP Conf. Series (EES) style via python-docx.

Mirrors the final manuscript in Paper-Tofu-Final-IOP-EES.md.

Run:
    python3 generate_paper_docx.py
"""

from __future__ import annotations

OUT = "/projects/sandbox/sonet-product/research-map/Paper-Tofu-Final-IOP-EES.docx"

TITLE = ("Life Cycle Assessment of Small-Scale Tofu Production in Semarang City: "
         "Hotspot Analysis Using CML-IA and AWARE Methods")

# Authors: list of (name, superscript)
AUTHORS = [("Abdul Jabbar", "1"), ("Putri Alifa Kholil", "1,*"),
           ("Desiana Fitri Awati", "1")]

AFFIL = ("1 Study Program of Environmental Science, Faculty of Mathematics and "
         "Natural Sciences, Universitas Negeri Semarang, Semarang, Indonesia")
EMAIL = "* E-mail: putrialifa@mail.unnes.ac.id"

ABSTRACT = (
    "The tofu industry is a vital food-processing sector in Indonesia, yet its production is "
    "water- and energy-intensive and generates high-strength organic wastewater. This study "
    "evaluated the potential environmental impacts of small-scale tofu production in Semarang "
    "City using a gate-to-gate Life Cycle Assessment (LCA). The system boundary covered washing, "
    "soaking, grinding, boiling, filtration, coagulation, and molding and packaging. Impact "
    "assessment was performed in SimaPro 10.3 using the CML-IA Baseline method for five "
    "categories\u2014Global Warming Potential (GWP), Ozone Layer Depletion (ODP), Freshwater "
    "Aquatic Ecotoxicity (FAET), Acidification (AP), and Eutrophication (EP)\u2014and water "
    "scarcity was assessed with the AWARE method. The washing and boiling stages were the "
    "dominant environmental hotspots. Total GWP reached 5.28 \u00d7 10\u00b3 kg CO\u2082 eq and "
    "FAET 1.33 \u00d7 10\u00b3 kg 1,4-DB eq, with washing contributing about 67% and 92% of these "
    "totals, respectively. After normalization, FAET showed the highest value "
    "(2.57 \u00d7 10\u207b\u2079), identifying organic liquid-waste management as the foremost "
    "environmental priority. The AWARE assessment yielded a total water scarcity of 515 m\u00b3 "
    "world eq, dominated by the boiling stage (57%) through the embedded water footprint of "
    "firewood. These results indicate that freshwater pollution, rather than climate impact "
    "alone, is the most acute burden at the facility scale. Recommended interventions\u2014"
    "anaerobic wastewater treatment with biogas recovery, water recycling, fuel substitution, "
    "and improved thermal efficiency\u2014target pollution at the source rather than end-of-pipe. "
    "The study provides process-level evidence to support cleaner production in urban tofu SMEs."
)

KEYWORDS = ("Life Cycle Assessment; Tofu Production; CML-IA Baseline; AWARE; "
            "Environmental Hotspot")

# Body blocks: each block is a dict describing its type.
# types: h1, h2, body, body_lead (lead-in bold + rest), caption, figure
BODY = [
    {"t": "h1", "text": "1. Introduction"},
    {"t": "body", "text":
        "Tofu (tahu) is a staple plant-based protein in Indonesia, consumed at roughly 0.16 kg "
        "per capita per week and produced largely by household-scale and small and medium-sized "
        "enterprises (SMEs) [1]. This sustained demand generates large volumes of high-strength "
        "organic wastewater rich in dissolved proteins, suspended solids, and soybean residues, "
        "which are frequently discharged with little or no treatment [2, 3]. Such effluents are a "
        "recognized driver of freshwater degradation: the food-processing sector is among the "
        "leading contributors to surface-water pollution in Indonesia, with untreated organic "
        "discharge elevating biochemical oxygen demand (BOD) and chemical oxygen demand (COD) in "
        "receiving rivers well beyond regulatory limits [4]. In dense urban settings such as "
        "Semarang City, where tofu SMEs typically operate close to riverine systems with limited "
        "treatment infrastructure, the ecological significance of process-level liquid waste is "
        "amplified [3]."},
    {"t": "body", "text":
        "Life Cycle Assessment (LCA), standardized under the ISO 14040 and ISO 14044 framework, "
        "is a rigorous approach for quantifying environmental burdens across a defined production "
        "system and for identifying the stages that disproportionately drive impacts [5]. LCA has "
        "been applied to Indonesian tofu production, but with recurring methodological gaps. "
        "Rosyidah et al. [6] and Nugroho et al. [7] focused mainly on greenhouse-gas emissions and "
        "identified firewood-fired boiling as the dominant emission source; Sari et al. [8] "
        "applied CML-IA in a cradle-to-gate study but did not resolve freshwater ecotoxicity at "
        "the stage level; and Hartini et al. [9] coupled LCA with life cycle costing and confirmed "
        "freshwater ecotoxicity and eutrophication as leading categories. However, none of these "
        "gate-to-gate Indonesian studies combined the CML-IA Baseline method with a dedicated, "
        "stage-resolved water-scarcity assessment [10]. Critically, the AWARE water-scarcity "
        "method has not been applied to Indonesian tofu production, and no LCA has been reported "
        "for Semarang City, where urban water dynamics may yield burden profiles distinct from "
        "rural production sites [11]."},
    {"t": "body", "text":
        "This study addresses these gaps by conducting a gate-to-gate LCA of SME-scale tofu "
        "production in Semarang City with four objectives: (1) to quantify impacts across five "
        "CML-IA Baseline midpoint categories\u2014GWP, ODP, FAET, AP, and EP\u2014using "
        "SimaPro 10.3; (2) to assess water scarcity with the AWARE method, the first such "
        "application to Indonesian tofu production [11]; (3) to characterize stage-level material "
        "and energy flows following standardized life cycle inventory guidance [12]; and (4) to "
        "identify environmental hotspots and propose targeted, source-oriented improvement "
        "strategies for urban tofu SMEs, consistent with cleaner-production and circular-economy "
        "principles [13]."},

    {"t": "h1", "text": "2. Methods"},
    {"t": "body", "text":
        "This research used a quantitative Life Cycle Assessment (LCA) approach to evaluate the "
        "environmental impacts of tofu production at an SME in Semarang City. The assessment "
        "followed SNI ISO 14040:2016 and SNI ISO 14044:2016 and comprised four sequential phases: "
        "(1) goal and scope definition, (2) life cycle inventory analysis, (3) life cycle impact "
        "assessment, and (4) interpretation [5]. Primary data were obtained through direct field "
        "measurements, structured interviews with the facility operator, and on-site documentation "
        "during a complete production batch. Secondary data were sourced from peer-reviewed "
        "literature and the Ecoinvent v3 database within SimaPro 10.3 [5]."},
    {"t": "h2", "text": "2.1 Goal and scope definition"},
    {"t": "body", "text":
        "The study assesses the environmental impacts of tofu production at the facility level, "
        "focusing on water use, energy consumption, and waste generation at each production stage. "
        "The functional unit was defined as one complete production batch (one production day), "
        "with 1,400 kg of soybeans as raw-material input yielding 3,681 kg of tofu; results are "
        "reported per batch to reflect the directly observed input\u2013output inventory, "
        "consistent with practice in site-specific SME-level assessments [8]. The system boundary "
        "adopts a gate-to-gate approach encompassing all unit processes from the initial washing "
        "of soybeans to final molding and packaging (Figure 1). Upstream operations (soybean "
        "farming and transportation) and downstream activities (distribution, consumption, and "
        "end-of-life) were excluded, consistent with prior gate-to-gate LCA studies of Indonesian "
        "tofu SMEs [8, 10]."},
    {"t": "caption", "text":
        "Figure 1. Gate-to-gate system boundary of the tofu production life cycle assessment."},
    {"t": "h2", "text": "2.2 Life cycle inventory analysis"},
    {"t": "body", "text":
        "The inventory phase identified and quantified all input and output flows for each unit "
        "process relative to the functional unit [12]. Data collected include soybean inputs, "
        "water and energy consumption, fuel use, and waste outputs comprising liquid waste, solid "
        "waste, and atmospheric emissions. The inventory data are presented in Table 1."},
    {"t": "h2", "text": "2.3 Impact assessment"},
    {"t": "body", "text":
        "The life cycle impact assessment (LCIA) phase evaluated the potential environmental "
        "impacts based on the compiled inventory [5]. The CML-IA Baseline method was applied in "
        "SimaPro 10.3 across five impact categories: Global Warming Potential (GWP\u2081\u2080\u2080\u2090, "
        "kg CO\u2082 eq), Ozone Layer Depletion (ODP, kg CFC-11 eq), Freshwater Aquatic "
        "Ecotoxicity (FAET, kg 1,4-DB eq), Acidification (AP, kg SO\u2082 eq), and Eutrophication "
        "(EP, kg PO\u2084\u00b3\u207b eq). Water use and potential water scarcity were additionally "
        "assessed using the AWARE (Available WAter REmaining) method, which weights water "
        "consumption against regional water availability and therefore captures water deprivation "
        "beyond simple volumetric accounting [11]."},
    {"t": "h2", "text": "2.4 Interpretation"},
    {"t": "body", "text":
        "The interpretation phase contextualized the LCIA results to identify hotspot categories "
        "and the production stages contributing the greatest burdens. Normalization was applied "
        "using the CML-IA Baseline reference values to express impact categories of different "
        "units on a common dimensionless scale, enabling their relative significance to be "
        "compared. Interpretation was confined to the gate-to-gate system boundary so that the "
        "results accurately reflect the stages exerting the greatest environmental influence [5]."},

    {"t": "h1", "text": "3. Results and Discussion"},
    {"t": "body", "text":
        "Tofu production within the gate-to-gate boundary involves a sequence of interdependent "
        "unit processes\u2014from soybean washing to molding and packaging\u2014each mobilizing "
        "distinct material and energy inputs and generating characteristic waste streams. As "
        "established in the inventory (Table 1), the system is dominated by high volumetric water "
        "consumption across multiple stages and by intensive thermal-energy demand during boiling, "
        "with firewood as the primary fuel. These characteristics create a multi-dimensional "
        "burden profile extending beyond greenhouse-gas emissions to freshwater pollution, "
        "nutrient loading, and acidification."},
    {"t": "h2", "text": "3.1 Potential environmental impacts of tofu production"},
    {"t": "body", "text":
        "The LCIA results per production stage are summarized in Table 2. A notable feature is the "
        "absence of impact values for the filtration stage across all categories. This is "
        "attributable to the absence of explicit energy inputs or additional material flows for "
        "that stage\u2014filtration operates as a passive mechanical separation step without "
        "recorded electricity consumption\u2014so SimaPro 10.3 did not generate characterization "
        "values for it. This is consistent with the inventory in Table 1, where filtration is "
        "limited to the passage of boiled slurry and the separation of soy milk from okara "
        "residue."},
    {"t": "body_lead", "lead": "Global Warming Potential (GWP).", "text":
        " Total GWP is 5.28 \u00d7 10\u00b3 kg CO\u2082 eq. The washing stage is the largest "
        "contributor at 3.56 \u00d7 10\u00b3 kg CO\u2082 eq (\u224867.4%), followed by boiling at "
        "1.63 \u00d7 10\u00b3 kg CO\u2082 eq (\u224830.9%). The dominance of washing is "
        "mechanistically driven by the large volume of water consumed (2,800 L per batch; "
        "Table 1), which is linked within the LCA framework to energy-intensive upstream water "
        "extraction, treatment, and distribution that generate indirect emissions [5]. In "
        "addition, organic-laden washing wastewater can generate methane (CH\u2084) under "
        "anaerobic decomposition if discharged untreated [2]. Boiling contributes substantially "
        "through the direct combustion of 800 kg of firewood per batch, releasing CO\u2082, "
        "CH\u2084, and N\u2082O\u2014consistent with comparative LCA studies of wood-fired tofu "
        "production in Indonesia, where boiling was identified as the primary emission source, "
        "accounting for up to ~98% of total CO\u2082-equivalent emissions in some production "
        "configurations [6]."},
    {"t": "body_lead", "lead": "Ozone Layer Depletion (ODP).", "text":
        " Total ODP is 6.11 \u00d7 10\u207b\u2075 kg CFC-11 eq, with washing contributing "
        "5.93 \u00d7 10\u207b\u2075 kg CFC-11 eq (\u224897.1%). This result may appear "
        "counterintuitive, as washing involves no halogenated substances; however, within the LCA "
        "framework the burden is allocated to the upstream energy system associated with "
        "water-supply infrastructure, where indirect emissions of trace ozone-depleting substances "
        "are attributed to processes consuming large volumes of grid-distributed water [5]. The "
        "disproportionate concentration of ODP in the washing stage therefore reflects the "
        "indirect burden of high water use."},
    {"t": "body_lead", "lead": "Freshwater Aquatic Ecotoxicity (FAET).", "text":
        " Total FAET is 1.33 \u00d7 10\u00b3 kg 1,4-DB eq, with washing contributing "
        "1.22 \u00d7 10\u00b3 kg 1,4-DB eq (\u224891.7%), followed by boiling (70.6 kg 1,4-DB eq; "
        "\u22485.3%) and grinding (24.6 kg 1,4-DB eq). The dominance of washing reflects the large "
        "effluent volume generated (2,660 L per batch), which contains residual organic matter, "
        "suspended solids, and impurities stripped from raw soybeans; in the CML-IA "
        "characterization, greater effluent volumes correspond to proportionally larger FAET "
        "values through fate-and-exposure modeling [9]. This aligns with prior Indonesian tofu LCA "
        "studies that consistently rank freshwater ecotoxicity among the three most significant "
        "categories [9]. The secondary contribution of boiling arises from atmospheric deposition "
        "of combustion-derived NO\u2093 and SO\u2093 into aquatic systems [6]."},
    {"t": "body_lead", "lead": "Acidification (AP).", "text":
        " Total AP is 4.59 kg SO\u2082 eq, dominated by washing (3.53 kg SO\u2082 eq; \u224876.9%), "
        "with boiling as the secondary contributor (0.723 kg SO\u2082 eq; \u224815.8%). The "
        "washing burden derives principally from the upstream energy system required for water "
        "provision, which emits SO\u2082 and NO\u2093\u2014the primary acidification precursors "
        "[5]. Firewood combustion during boiling adds acidifying emissions that may form acid "
        "deposition, with potential consequences for soil and receiving waters near the facility."},
    {"t": "body_lead", "lead": "Eutrophication (EP).", "text":
        " Total EP is 9.10 kg PO\u2084\u00b3\u207b eq, with washing dominant at 8.30 kg "
        "PO\u2084\u00b3\u207b eq (\u224891.2%). This is driven by nutrient-rich effluents "
        "containing nitrogen and phosphorus compounds leached from soybeans; if discharged "
        "untreated, these can trigger excessive algal proliferation, dissolved-oxygen depletion, "
        "and disruption of aquatic ecosystems [2]. Boiling contributes 0.409 kg PO\u2084\u00b3\u207b "
        "eq (\u22484.5%) through atmospheric deposition of combustion-derived NO\u2093."},
    {"t": "h2", "text": "3.2 Water use assessment (AWARE)"},
    {"t": "body", "text":
        "Water use was assessed with the AWARE method, which expresses the volume of water "
        "deprived from downstream users per unit consumed at a given location. The total "
        "water-scarcity impact of the production system is 515 m\u00b3 world eq (Table 3). Contrary "
        "to the pattern in the CML-IA categories, the boiling stage is the dominant contributor at "
        "294 m\u00b3 world eq (\u224857.1%), rather than washing. This reflects the AWARE method's "
        "sensitivity to upstream supply-chain water consumption: firewood procurement, processing, "
        "and distribution embed substantial water use that is allocated back to boiling, elevating "
        "its scarcity value beyond what direct water consumption alone would suggest [11]. Grinding "
        "contributes the second-largest share (101 m\u00b3 world eq; \u224819.6%), consistent with "
        "the continuous water addition required to achieve appropriate slurry consistency "
        "(4,200 L; Table 1). The soaking, coagulation, and washing stages contribute 39.3, 38.2, "
        "and 35.8 m\u00b3 world eq (\u22487.6%, 7.4%, and 7.0%, respectively), and molding and "
        "packaging the least (7.63 m\u00b3 world eq; \u22481.5%). The comparatively lower scarcity "
        "value of washing\u2014despite generating the largest effluent volume\u2014illustrates that "
        "water-scarcity characterization is governed not solely by effluent volume but by regional "
        "water availability and the embedded water footprint of upstream inputs [11]."},
    {"t": "h2", "text": "3.3 Normalization"},
    {"t": "body", "text":
        "Normalization was performed using the CML-IA Baseline reference values to place all "
        "categories on a common dimensionless scale (Table 4). The results confirm that FAET is "
        "the dominant impact category (2.57 \u00d7 10\u207b\u2079), identifying tofu production as "
        "exerting its greatest relative pressure on freshwater ecosystems. This is significant "
        "because FAET ranks second in absolute terms behind GWP, yet emerges as dominant after "
        "normalization\u2014indicating that the ecotoxicity burden is disproportionately large "
        "relative to global reference values [9]. The finding affirms that organic liquid-waste "
        "streams\u2014particularly washing effluents and coagulation whey\u2014are the most "
        "ecologically consequential outputs of this system."},
    {"t": "body", "text":
        "Although GWP frequently receives primary attention in sustainability discourse, its "
        "normalized value ranks second (1.05 \u00d7 10\u207b\u2079), indicating that freshwater "
        "pollution is a more acute concern at the facility scale assessed [6]. Eutrophication "
        "(6.89 \u00d7 10\u207b\u00b9\u2070) reinforces the urgency of addressing nutrient-laden "
        "wastewater, whereas Acidification (1.63 \u00d7 10\u207b\u00b9\u2070) and ODP "
        "(6.84 \u00d7 10\u207b\u00b9\u00b3) contribute relatively minor normalized values and are "
        "therefore not the primary environmental priorities for this system."},
    {"t": "h2", "text": "3.4 Hotspot identification and improvement strategies"},
    {"t": "body", "text":
        "The integrated analysis of LCIA and normalization results identifies two primary "
        "hotspots: (i) organic liquid-waste generation across the washing, soaking, and "
        "coagulation stages, and (ii) firewood combustion during boiling. These correspond, "
        "respectively, to the dominant drivers of FAET and EP (ranked first and third in "
        "normalized impact) and of GWP and water scarcity. The recommended, evidence-based "
        "improvement strategies are summarized in Table 5."},
    {"t": "body", "text":
        "For the first hotspot, the high-BOD/COD, nutrient-rich effluents from washing, soaking, "
        "and coagulation pose significant risks to receiving ecosystems if discharged untreated "
        "[2, 3]. Anaerobic treatment (e.g., biodigesters or integrated biogas wastewater units) "
        "is technically appropriate for SME-scale operations because it reduces organic loading "
        "while recovering biogas as an energy resource, simultaneously lowering FAET (through "
        "effluent-load reduction) and GWP (through fuel displacement) [2]. Complementary "
        "measures include constructed-wetland or biofilter polishing for nutrient removal, "
        "counter-current water recycling within washing and soaking to cut both water consumption "
        "and effluent generation, and valorization of whey as a protein-rich feed or fermentation "
        "substrate to divert high-nutrient streams away from treatment [13]."},
    {"t": "body", "text":
        "For the second hotspot, boiling contributes disproportionately to GWP and water scarcity "
        "owing to its reliance on firewood: combustion of 800 kg per batch releases 1,397 kg "
        "CO\u2082, 0.0499 kg CH\u2084, and 0.3744 kg N\u2082O. Transitioning to cleaner energy\u2014"
        "such as LPG or biomethane recovered from on-site biodigesters\u2014would directly reduce "
        "these emissions while lowering the upstream water footprint embedded in the firewood "
        "supply chain [14]. Adopting improved-combustion or thermally insulated boiling systems, "
        "together with waste-heat recovery to pre-heat process water, would further reduce specific "
        "fuel consumption [14]. Overall, the LCA evidence demonstrates that interventions targeting "
        "liquid-waste management and energy efficiency at the source yield substantially greater "
        "environmental benefit than conventional end-of-pipe approaches [13]."},
    {"t": "body", "text":
        "Taken together, the disproportionate contribution of the washing stage to FAET, GWP, "
        "acidification, and eutrophication\u2014evidenced consistently across Tables 2 and 4 and "
        "Figure 2\u2014indicates that process-level modifications at the earliest stages of the "
        "production chain deliver the greatest systemic environmental benefit. This conclusion is "
        "consistent with the waste-reduction hierarchy embedded in the ISO 14040/14044 framework "
        "[5] and with circular-production principles that prioritize source reduction and resource "
        "recovery over end-of-pipe control [13]."},

    {"t": "h1", "text": "4. Conclusion"},
    {"t": "body", "text":
        "This study applied a gate-to-gate Life Cycle Assessment to evaluate the environmental "
        "impacts of small-scale tofu production in Semarang City using the CML-IA Baseline and "
        "AWARE methods. Tofu production generated impacts in all assessed categories, with totals "
        "of 5.28 \u00d7 10\u00b3 kg CO\u2082 eq (GWP), 6.11 \u00d7 10\u207b\u2075 kg CFC-11 eq "
        "(ODP), 1.33 \u00d7 10\u00b3 kg 1,4-DB eq (FAET), 4.59 kg SO\u2082 eq (AP), and 9.10 kg "
        "PO\u2084\u00b3\u207b eq (EP). The washing and boiling stages were the principal hotspots: "
        "washing dominated freshwater aquatic ecotoxicity and eutrophication through "
        "organic-wastewater generation, while boiling drove GWP and water scarcity through "
        "firewood combustion and its embedded water demand. Normalization identified FAET as the "
        "most significant category (2.57 \u00d7 10\u207b\u2079), and the AWARE assessment yielded a "
        "total water-scarcity impact of 515 m\u00b3 world eq, dominated by boiling."},
    {"t": "body", "text":
        "These findings indicate that freshwater pollution\u2014rather than climate impact "
        "alone\u2014is the most acute environmental burden of urban tofu SMEs, and that "
        "environmental management should prioritize organic liquid-waste treatment, water reuse, "
        "and boiling-stage energy efficiency. Practical, high-leverage measures include anaerobic "
        "treatment with biogas recovery, counter-current water recycling, fuel substitution, and "
        "improved thermal systems, all of which act upstream at the source rather than "
        "end-of-pipe. By providing the first stage-resolved, AWARE-supported LCA of tofu "
        "production in Semarang City, this study contributes process-level evidence to support "
        "cleaner production in urban food SMEs. Future research should extend the boundary toward "
        "cradle-to-grave and integrate life cycle costing to confirm the economic feasibility of "
        "the proposed interventions."},
]

# Tables (caption + rows; first row is header)
TABLE1_CAP = "Table 1. Life cycle inventory of tofu production (per batch / production day)."
TABLE1 = [
    ["Process stage", "Inputs", "Quantity", "Unit", "Outputs", "Quantity", "Unit"],
    ["Washing", "Soybeans", "1,400", "kg", "Washed soybeans", "1,540", "kg"],
    ["", "Water", "2,800", "L", "Wastewater", "2,660", "L"],
    ["", "Electricity", "0.60", "kWh", "CO\u2082 emissions", "0.51", "kg"],
    ["Soaking", "Washed soybeans", "1,540", "kg", "Hydrated soybeans", "3,640", "kg"],
    ["", "Water", "2,700", "L", "Wastewater", "600", "L"],
    ["", "Electricity", "0.45", "kWh", "CO\u2082 emissions", "0.38", "kg"],
    ["Grinding", "Hydrated soybeans", "3,640", "kg", "Soybean slurry", "7,840", "kg"],
    ["", "Water", "4,200", "L", "CO\u2082 emissions", "14.45", "kg"],
    ["", "Electricity", "17", "kWh", "", "", ""],
    ["Boiling", "Soybean slurry", "7,840", "kg", "Boiled slurry", "15,652", "kg"],
    ["", "Water", "8,400", "L", "Steam loss (evaporation)", "588", "L"],
    ["", "Firewood", "800", "kg", "CO\u2082 emissions", "1,397", "kg"],
    ["", "", "", "", "CH\u2084 emissions", "0.0499", "kg"],
    ["", "", "", "", "N\u2082O emissions", "0.3744", "kg"],
    ["Filtration", "Boiled slurry", "15,652", "kg", "Soy milk", "14,052", "kg"],
    ["", "", "", "", "Okara (residue)", "1,600", "kg"],
    ["Coagulation", "Soy milk", "14,052", "kg", "Tofu curd", "3,681", "kg"],
    ["", "Coagulant water", "421", "L", "Whey", "10,434", "L"],
    ["Molding & packaging", "Tofu curd", "3,681", "kg", "Final tofu product", "3,681", "kg"],
    ["", "Water", "200", "L", "Wastewater", "100", "L"],
]

TABLE2_CAP = "Table 2. Environmental impact of tofu production by stage."
TABLE2 = [
    ["Impact category", "Unit", "Total", "Washing", "Soaking", "Grinding", "Boiling",
     "Coagulation", "Molding & packing"],
    ["Global Warming (GWP\u2081\u2080\u2080\u2090)", "kg CO\u2082 eq", "5.28 \u00d7 10\u00b3",
     "3.56 \u00d7 10\u00b3", "0.945", "35.8", "1.63 \u00d7 10\u00b3", "55.8", "5.9"],
    ["Ozone Layer Depletion (ODP)", "kg CFC-11 eq", "6.11 \u00d7 10\u207b\u2075",
     "5.93 \u00d7 10\u207b\u2075", "2.10 \u00d7 10\u207b\u2079", "7.93 \u00d7 10\u207b\u2078",
     "1.07 \u00d7 10\u207b\u2076", "5.02 \u00d7 10\u207b\u2077", "1.29 \u00d7 10\u207b\u2077"],
    ["Freshwater Aquatic Ecotoxicity (FAET)", "kg 1,4-DB eq", "1.33 \u00d7 10\u00b3",
     "1.22 \u00d7 10\u00b3", "0.651", "24.6", "70.6", "14.1", "2.75"],
    ["Acidification (AP)", "kg SO\u2082 eq", "4.59", "3.53", "0.00245", "0.0925", "0.723",
     "0.226", "0.0176"],
    ["Eutrophication (EP)", "kg PO\u2084\u00b3\u207b eq", "9.10", "8.30", "0.0034", "0.128",
     "0.409", "0.25", "0.00838"],
]

TABLE3_CAP = "Table 3. Water use assessment (AWARE) results of tofu production."
TABLE3 = [
    ["Production stage", "Water scarcity (m\u00b3 world eq)", "Share (%)"],
    ["Boiling", "294", "57.09"],
    ["Grinding", "101", "19.61"],
    ["Soaking", "39.3", "7.63"],
    ["Coagulation", "38.2", "7.42"],
    ["Washing", "35.8", "6.95"],
    ["Molding & packaging", "7.63", "1.48"],
    ["Total", "515", "100"],
]

TABLE4_CAP = "Table 4. Normalization results of tofu production environmental impacts."
TABLE4 = [
    ["Impact category", "Normalized value"],
    ["Freshwater Aquatic Ecotoxicity (FAET)", "2.57 \u00d7 10\u207b\u2079"],
    ["Global Warming Potential (GWP\u2081\u2080\u2080\u2090)", "1.05 \u00d7 10\u207b\u2079"],
    ["Eutrophication (EP)", "6.89 \u00d7 10\u207b\u00b9\u2070"],
    ["Acidification (AP)", "1.63 \u00d7 10\u207b\u00b9\u2070"],
    ["Ozone Layer Depletion (ODP)", "6.84 \u00d7 10\u207b\u00b9\u00b3"],
]

TABLE5_CAP = "Table 5. Recommended environmental improvement strategies for tofu production."
TABLE5 = [
    ["Improvement strategy", "Description", "Reference"],
    ["Anaerobic wastewater treatment / biofilter system",
     "Reduces organic load (BOD, COD) from washing, soaking, and whey effluents; anaerobic "
     "digestion also recovers biogas as a renewable energy source, providing co-benefits for "
     "FAET and GWP.", "[2]"],
    ["Water recycling (counter-current washing)",
     "Reuses water at the washing and soaking stages to lower clean-water consumption and "
     "minimize effluent volume and organic concentration.", "[3]"],
    ["Fuel substitution and thermal-efficiency improvement at boiling",
     "Replaces firewood with LPG or biomethane (recovered from tofu wastewater) and adopts "
     "insulated/improved boilers and waste-heat recovery, reducing combustion emissions and the "
     "embedded water footprint.", "[14]"],
]

FIG2_CAP = ("Figure 2. Environmental impact of tofu production across production stages "
            "(SimaPro 10.3 output).")

REFERENCES = [
    "[1] Badan Pusat Statistik 2026 Rata-Rata Konsumsi per Kapita Seminggu Beberapa Macam Bahan "
    "Makanan Penting 2007\u20132025 (Jakarta: Badan Pusat Statistik) Available at: "
    "https://www.bps.go.id",
    "[2] Seroja R, Effendi H and Hariyadi S 2018 Tofu wastewater treatment using vetiver grass "
    "(Vetiveria zizanioides) and zeliac Appl. Water Sci. 8 2 "
    "https://doi.org/10.1007/s13201-018-0640-y",
    "[3] Sjafruddin R, Agustang A and Pertiwi N 2022 Estimasi limbah industri tahu dan kajian "
    "penerapan sistem produksi bersih J. Ilm. Mandala Educ. 8 1229\u201337 "
    "https://doi.org/10.36312/jime.v8i2.2826",
    "[4] Basuki T M, Nugroho H Y S H, Indrajaya Y, Pramono I B, Nugroho N P, Supangat A B et al. "
    "2024 Water pollution of some major rivers in Indonesia: status, institution, regulation, and "
    "recommendation for mitigation Pol. J. Environ. Stud. 33 3515\u201330 "
    "https://doi.org/10.15244/pjoes/178532",
    "[5] Chitaka T Y and Goga T 2023 The evolution of life cycle assessment in the food and "
    "beverage industry: a review Cambridge Prisms: Plast. 1 e2 https://doi.org/10.1017/plc.2023.4",
    "[6] Rosyidah M, Masruri A and Putra R A 2020 Analysis of environmental impact with the Life "
    "Cycle Assessment (LCA) method on tofu production Int. J. Sci. Technol. Manag. 1 428\u201335 "
    "https://doi.org/10.46729/ijstm.v1i4.73",
    "[7] Nugroho M E, Setyono P and Rachmawati S 2024 Analisis emisi gas rumah kaca dengan Life "
    "Cycle Assessment (LCA) dan Analytical Hierarchy Process (AHP) industri tahu J. Ilmu Lingkung. "
    "22 1504\u201312 https://doi.org/10.14710/jil.22.6.1504-1512",
    "[8] Sari I P, Kurniawan W and Sia F L 2021 Environmental impact of tofu production in West "
    "Jakarta using a life cycle assessment approach IOP Conf. Ser.: Earth Environ. Sci. 896 012050 "
    "https://doi.org/10.1088/1755-1315/896/1/012050",
    "[9] Hartini S, Fatliana A N, Handayani N U, Wicaksono P A, Ramadan B S and Matsumoto T 2024 "
    "Life cycle assessment and life cycle cost of tofu production and its extended recycling "
    "scenario Glob. J. Environ. Sci. Manag. 10 487 "
    "[DOI: add from published version, prefix 10.22034/gjesm]",
    "[10] Kartika Wardana S, Cucikodana Y, Almaniar S, Dwijayanti A, Maulana F and Muhlisoh N A "
    "2024 Penilaian dampak lingkungan dengan menggunakan Life Cycle Assessment (LCA) pada industri "
    "tahu Kampung Jangkar Kulon, Cilegon Banten J. Teknol. Kim. Unimal 13 97\u2013106 "
    "https://doi.org/10.29103/jtku.v13i2.16429",
    "[11] Mir B A, Nurdiawati A and Al-Ghamdi S G 2025 Assessing the environmental impact of "
    "freshwater use in LCA: established practices and current methods Environ. Sci.: Water Res. "
    "Technol. 11 196\u2013221 https://doi.org/10.1039/D4EW00641K",
    "[12] Saavedra-Rubio K, Thonemann N, Crenna E, Lemoine B, Caliandro P and Laurent A 2022 "
    "Stepwise guidance for data collection in the life cycle inventory (LCI) phase: building "
    "technology-related LCI blocks J. Clean. Prod. 366 132903 "
    "https://doi.org/10.1016/j.jclepro.2022.132903",
    "[13] Bj\u00f8rnbet M M and Vild\u00e5sen S S 2021 Life cycle assessment to ensure "
    "sustainability of circular business models in manufacturing Sustainability 13 11014 "
    "https://doi.org/10.3390/su131911014",
    "[14] Ningsih L M, Mazancov\u00e1 J, Hasanudin U and Roub\u00edk H 2026 Energy audits in the "
    "tofu industry: an evaluation of energy consumption towards a green and sustainable industry "
    "Environ. Dev. Sustain. 28 3737\u201359 https://doi.org/10.1007/s10668-024-05109-z",
]


def build_docx(out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    base = doc.styles["Normal"]
    base.font.name = "Times New Roman"
    base.font.size = Pt(11)

    def _set(p, *, align=None, before=0, after=6, line=1.15, indent=None):
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line
        if align is not None:
            p.alignment = align
        if indent is not None:
            pf.first_line_indent = Cm(indent)

    def para(runs, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6,
             line=1.15, size=11, indent=None):
        """runs: str or list of (text, bold, italic, superscript, subscript)."""
        p = doc.add_paragraph()
        if isinstance(runs, str):
            runs = [(runs, False, False, False, False)]
        for item in runs:
            text, bold, italic, sup, sub = item
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.bold = bold
            r.italic = italic
            r.font.superscript = sup
            r.font.subscript = sub
        _set(p, align=align, before=before, after=after, line=line, indent=indent)
        return p

    def set_cell_borders(cell):
        tcpr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "4")
            e.set(qn("w:color"), "000000")
            borders.append(e)
        tcpr.append(borders)

    def shade(cell, color="D9E2F3"):
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tcpr.append(shd)

    def add_table(rows, *, font_size=8.5, left_cols=(0,)):
        ncols = len(rows[0])
        tbl = doc.add_table(rows=len(rows), cols=ncols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = tbl.cell(i, j)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_borders(cell)
                if i == 0:
                    shade(cell)
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(str(val))
                r.font.name = "Times New Roman"
                r.font.size = Pt(font_size)
                if i == 0:
                    r.bold = True
                pf = p.paragraph_format
                pf.space_before = Pt(1)
                pf.space_after = Pt(1)
                pf.line_spacing = 1.0
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif j in left_cols:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return tbl

    def caption(text):
        para([(text, False, False, False, False)],
             align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4, size=9.5)

    # ---- Title ----
    para([(TITLE, True, False, False, False)],
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10, size=15, line=1.15)

    # ---- Authors ----
    author_runs = []
    for idx, (name, sup) in enumerate(AUTHORS):
        author_runs.append((name, False, False, False, False))
        author_runs.append((sup, False, False, True, False))
        if idx < len(AUTHORS) - 1:
            author_runs.append((", ", False, False, False, False))
    para(author_runs, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4, size=11)

    # ---- Affiliation ----
    para([(AFFIL[0], False, False, True, False),
          (AFFIL[1:], False, True, False, False)],
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, size=9.5)
    para([(EMAIL, False, True, False, False)],
         align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, size=9.5)

    # ---- Abstract ----
    para([("Abstract. ", True, False, False, False),
          (ABSTRACT, False, False, False, False)],
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, size=10,
         indent=None)
    para([("Keywords: ", True, False, False, False),
          (KEYWORDS, False, False, False, False)],
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=12, size=10)

    # ---- Body ----
    def emit_body_text(text, size=11):
        """Render body text, converting standalone sub/superscript markers handled inline.
        Here text is plain (Unicode already), so single run."""
        para([(text, False, False, False, False)],
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, size=size)

    for blk in BODY:
        t = blk["t"]
        if t == "h1":
            para([(blk["text"], True, False, False, False)],
                 align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=5, size=12)
        elif t == "h2":
            para([(blk["text"], True, True, False, False)],
                 align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4, size=11)
        elif t == "body":
            emit_body_text(blk["text"])
        elif t == "body_lead":
            para([(blk["lead"], True, False, False, False),
                  (blk["text"], False, False, False, False)],
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, size=11)
        elif t == "caption":
            caption(blk["text"])

    # ---- Tables section ----
    para([("Tables", True, False, False, False)],
         align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=5, size=12)

    caption(TABLE1_CAP); add_table(TABLE1, left_cols=(0, 1, 4))
    caption(TABLE2_CAP); add_table(TABLE2, left_cols=(0,))
    caption(TABLE3_CAP); add_table(TABLE3, left_cols=(0,))
    caption(TABLE4_CAP); add_table(TABLE4, left_cols=(0,))
    caption(TABLE5_CAP); add_table(TABLE5, font_size=9, left_cols=(0, 1))
    caption(FIG2_CAP)

    # ---- References ----
    para([("References", True, False, False, False)],
         align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=5, size=12)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.0
        pf.left_indent = Cm(0.75)
        pf.first_line_indent = Cm(-0.75)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == "__main__":
    build_docx(OUT)
