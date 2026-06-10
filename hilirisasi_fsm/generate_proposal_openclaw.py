"""
Generator: Proposal Hilirisasi Riset OpenClaw.

Menduplikasi FORMAT dari:
    Proposal-hilirisasi-riset-kinerja-pegawai-v1.2.docx
dan hanya mengganti KONTEN menjadi konteks:
    "Pengembangan Sistem Monitoring dan Otomatisasi Web dan Media Sosial
     Terintegrasi berbasis Agentic AI untuk Mendukung Promosi Digital
     Fakultas Sains dan Matematika Universitas Diponegoro"
dengan Agentic AI utama: OpenClaw.

Output:
    Proposal-hilirisasi-riset-openclaw-v1.docx

Catatan:
    - Format (tata letak, font, shading section A-H, gambar, tabel bersarang,
      halaman pengesahan, dll.) DIPERTAHANKAN apa adanya dari dokumen sumber.
    - Hanya teks substansi (judul, ringkasan, kata kunci, pendahuluan, metode,
      tanggung jawab tim, jadwal, dan daftar pustaka) yang diganti.

Run:
    python3 generate_proposal_openclaw.py
"""

from __future__ import annotations

import copy
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "Proposal-hilirisasi-riset-kinerja-pegawai-v1.2.docx")
OUT = os.path.join(HERE, "Proposal-hilirisasi-riset-openclaw-v1.docx")

# =============================================================================
# KONTEN BARU
# =============================================================================

TITLE = ("Pengembangan Sistem Monitoring dan Otomatisasi Web dan Media Sosial "
         "Terintegrasi berbasis Agentic AI untuk Mendukung Promosi Digital "
         "Fakultas Sains dan Matematika Universitas Diponegoro")

KEYWORDS = ("Agentic AI; OpenClaw; Otomatisasi Media Sosial; Monitoring Web; "
            "Promosi Digital")

RINGKASAN = (
    "Promosi digital telah menjadi instrumen strategis bagi perguruan tinggi "
    "dalam membangun reputasi, menjaring calon mahasiswa, dan mendiseminasikan "
    "capaian akademik. Fakultas Sains dan Matematika (FSM) Universitas "
    "Diponegoro mengelola beragam kanal digital—situs web fakultas, situs "
    "departemen, serta akun media sosial seperti Instagram, X (Twitter), "
    "YouTube, dan LinkedIn. Namun pengelolaan kanal tersebut saat ini masih "
    "dilakukan secara manual, terfragmentasi, dan reaktif, sehingga "
    "pemantauan performa, penjadwalan konten, serta analisis keterlibatan "
    "(engagement) tidak terintegrasi dan sulit dilakukan secara konsisten. "
    "Penelitian ini bertujuan mengembangkan Sistem Monitoring dan Otomatisasi "
    "Web dan Media Sosial Terintegrasi berbasis Agentic AI dengan kerangka "
    "utama OpenClaw. Sistem dirancang sebagai orkestrasi beberapa agen "
    "otonom—agen pemantau (monitoring agent), agen otomatisasi konten "
    "(content automation agent), agen analitik (analytics agent), dan agen "
    "penjadwalan (scheduling agent)—yang berkolaborasi memantau kondisi "
    "situs web dan media sosial, menghasilkan serta menjadwalkan konten "
    "promosi, dan menyajikan analisis keterlibatan secara real-time pada "
    "satu dashboard. Penelitian menggunakan pendekatan Research and "
    "Development dengan model prototyping iteratif. Pengujian mencakup uji "
    "fungsional black-box, uji usability menggunakan System Usability Scale "
    "(SUS), serta uji performa sistem. Luaran yang ditargetkan berupa "
    "prototipe sistem promosi digital berbasis Agentic AI yang siap "
    "dioperasikan, publikasi ilmiah, dan pendaftaran Hak Kekayaan "
    "Intelektual (HKI). Sistem ini diharapkan meningkatkan efisiensi, "
    "konsistensi, dan efektivitas promosi digital FSM Undip secara "
    "terukur dan berbasis data."
)

# --- D. PENDAHULUAN ----------------------------------------------------------
# Tiap item: ("h", teks) heading tebal | ("p", teks) paragraf | ("b", "") baris kosong
PENDAHULUAN = [
    ("h", "Latar Belakang"),
    ("p", "Transformasi digital telah mengubah cara institusi pendidikan "
          "tinggi membangun reputasi dan menjalin komunikasi dengan publik. "
          "Kehadiran digital (digital presence) yang kuat melalui situs web "
          "dan media sosial kini menjadi faktor penentu daya saing fakultas, "
          "khususnya dalam menjaring calon mahasiswa, menarik mitra "
          "kerja sama, serta mendiseminasikan capaian riset dan akademik [6], "
          "[7]. Media sosial telah berkembang dari sekadar kanal komunikasi "
          "menjadi instrumen pemasaran strategis yang menuntut pengelolaan "
          "konten yang konsisten, terukur, dan responsif terhadap dinamika "
          "audiens [6], [8]."),
    ("p", "Fakultas Sains dan Matematika (FSM) Universitas Diponegoro "
          "mengelola beragam kanal digital, mulai dari situs web fakultas dan "
          "departemen hingga sejumlah akun media sosial seperti Instagram, X "
          "(Twitter), YouTube, dan LinkedIn. Namun, pengelolaan kanal-kanal "
          "tersebut saat ini masih dilakukan secara manual dan "
          "terfragmentasi: pemantauan performa situs, pembuatan dan "
          "penjadwalan konten, serta analisis keterlibatan audiens dikerjakan "
          "secara terpisah dan bergantung pada ketersediaan sumber daya "
          "manusia. Akibatnya, promosi digital menjadi reaktif, tidak "
          "konsisten, dan sulit dievaluasi dampaknya secara objektif."),
    ("b", ""),
    ("p", "Kemajuan kecerdasan buatan generatif dan model bahasa besar (Large "
          "Language Model/LLM) telah melahirkan paradigma baru berupa Agentic "
          "AI, yaitu sistem agen otonom yang mampu merencanakan, mengambil "
          "tindakan, menggunakan perkakas (tools), dan berkolaborasi untuk "
          "menyelesaikan tugas kompleks secara mandiri [1], [2], [3]. "
          "Pendekatan agen berbasis LLM seperti penalaran-dan-tindakan "
          "(reasoning-and-acting) terbukti efektif memadukan kemampuan "
          "bernalar dengan eksekusi tindakan nyata pada lingkungan digital "
          "[3], [4]. Paradigma ini sangat relevan untuk mengotomatiskan "
          "alur kerja promosi digital yang berulang namun membutuhkan "
          "adaptasi kontekstual."),
    ("p", "Pada penelitian ini, kerangka Agentic AI yang digunakan sebagai "
          "tulang punggung sistem adalah OpenClaw, sebuah kerangka kerja "
          "agen otonom bersifat terbuka yang memungkinkan orkestrasi "
          "multi-agen, integrasi perkakas eksternal (web crawler, "
          "Application Programming Interface/API media sosial, basis data), "
          "serta penjadwalan tugas secara terprogram. Melalui OpenClaw, "
          "beragam agen dapat dirancang untuk saling berkolaborasi: memantau "
          "kesehatan dan performa situs web, menganalisis tren keterlibatan "
          "media sosial, menyusun draf konten promosi, hingga menjadwalkan "
          "publikasi secara otomatis."),
    ("b", ""),
    ("h", "Rumusan Masalah"),
    ("p", "Bagaimana mengembangkan sistem monitoring dan otomatisasi web dan "
          "media sosial terintegrasi berbasis Agentic AI (OpenClaw) yang "
          "mampu memantau performa kanal digital, menghasilkan dan "
          "menjadwalkan konten promosi secara otomatis, serta menyajikan "
          "analitik keterlibatan secara real-time untuk mendukung promosi "
          "digital FSM Universitas Diponegoro secara efisien, konsisten, dan "
          "berbasis data?"),
    ("b", ""),
    ("h", "State-of-the-Art"),
    ("p", "Beberapa state-of-the-art yang melandasi penelitian ini adalah "
          "sebagai berikut:"),
    ("p", "Wang et al. [1] dan Xi et al. [2] melalui kajian sistematis "
          "menegaskan potensi besar agen otonom berbasis LLM dalam "
          "perencanaan, penggunaan perkakas, dan kolaborasi multi-agen untuk "
          "menyelesaikan tugas dunia nyata yang kompleks."),
    ("p", "Yao et al. [3] memperkenalkan kerangka penalaran-dan-tindakan "
          "(ReAct) yang memadukan rantai penalaran dengan aksi pada "
          "lingkungan eksternal, sementara Park et al. [4] menunjukkan agen "
          "generatif mampu menampilkan perilaku otonom yang koheren dan "
          "adaptif."),
    ("p", "Pada ranah pemasaran, Kaplan dan Haenlein [6] serta Peruta dan "
          "Shields [7] menggarisbawahi peran sentral media sosial dalam "
          "strategi komunikasi institusi pendidikan tinggi, dan Todor [9] "
          "menegaskan bahwa otomatisasi pemasaran meningkatkan efisiensi "
          "serta konsistensi aktivitas promosi."),
    ("b", ""),
    ("h", "Kebaruan"),
    ("p", "Kebaruan penelitian ini terletak pada:"),
    ("p", "1. Integrasi tunggal antara monitoring web, monitoring media "
          "sosial, otomatisasi pembuatan dan penjadwalan konten, serta "
          "analitik keterlibatan dalam satu sistem terorkestrasi berbasis "
          "Agentic AI untuk konteks promosi digital perguruan tinggi."),
    ("p", "2. Pemanfaatan kerangka Agentic AI OpenClaw sebagai orkestrator "
          "multi-agen otonom (monitoring, konten, analitik, dan penjadwalan) "
          "yang belum banyak didokumentasikan secara akademis untuk "
          "kebutuhan promosi digital institusional."),
    ("b", ""),
    ("h", "Peta Jalan Penelitian"),
    ("b", ""),
    ("IMG1", ""),
    ("CAP", "Gambar 1. Peta jalan penelitian selama 5 Tahun"),
    ("b", ""),
    ("p", "Penelitian ini merupakan bagian dari peta jalan penelitian jangka "
          "panjang selama lima tahun (2025\u20132029) di Departemen "
          "Informatika FSM UNDIP. Pada tahap awal difokuskan pada "
          "pengembangan fondasi sistem agentic untuk monitoring dan "
          "otomatisasi kanal digital, dilanjutkan dengan perluasan cakupan "
          "kanal dan integrasi analitik prediktif, hingga pada tahap akhir "
          "menjadi model rujukan promosi digital cerdas berbasis Agentic AI "
          "yang dapat direplikasi ke fakultas lain di Undip maupun perguruan "
          "tinggi sejawat."),
]

# --- E. METODE ---------------------------------------------------------------
METODE = [
    ("b", ""),
    ("p", "Penelitian ini menggunakan pendekatan Research and Development "
          "(R&D) dengan model pengembangan perangkat lunak berbasis "
          "prototyping iteratif. Pemilihan metode prototyping didasarkan pada "
          "kebutuhan untuk memvalidasi perilaku agen otonom secara bertahap "
          "bersama pemangku kepentingan promosi digital fakultas, serta untuk "
          "menyempurnakan orkestrasi agen secara berkelanjutan berdasarkan "
          "umpan balik nyata."),
    ("b", ""),
    ("IMG2", ""),
    ("CAP", "Gambar 2. Diagram alir penelitian"),
    ("CAPB", ""),
    ("p", "Tahapan penelitian terdiri dari enam tahap yang dilaksanakan "
          "secara iteratif. Tahap pertama adalah Analisis Kebutuhan dan "
          "Pemetaan Kanal Digital, yang mencakup inventarisasi situs web dan "
          "akun media sosial FSM, identifikasi indikator promosi "
          "(jangkauan, keterlibatan, pertumbuhan pengikut, performa situs), "
          "serta perumusan kebutuhan fungsional sistem dan peran tiap agen."),
    ("b", ""),
    ("p", "Tahap kedua adalah Perancangan Arsitektur dan Orkestrasi Agentic "
          "AI. Arsitektur sistem dirancang berbasis kerangka OpenClaw sebagai "
          "orkestrator multi-agen, dipadukan dengan modul web crawler untuk "
          "pemantauan situs, integrasi API media sosial (Instagram, X, "
          "YouTube, LinkedIn), basis data terpusat, serta dashboard "
          "visualisasi. Pada tahap ini didefinisikan peran agen pemantau "
          "(monitoring agent), agen otomatisasi konten (content automation "
          "agent), agen analitik (analytics agent), dan agen penjadwalan "
          "(scheduling agent)."),
    ("b", ""),
    ("p", "Tahap ketiga adalah Implementasi Iteratif yang berlangsung dalam "
          "tiga sprint. Sprint 1 mengembangkan monitoring agent untuk web dan "
          "media sosial beserta integrasi API. Sprint 2 mengembangkan content "
          "automation agent untuk generasi serta penjadwalan konten promosi "
          "berbantuan LLM. Sprint 3 mengembangkan analytics agent dan "
          "dashboard analitik keterlibatan, termasuk mekanisme rekomendasi "
          "strategi promosi."),
    ("b", ""),
    ("p", "Tahap keempat adalah Evaluasi dan Validasi Prototipe yang mencakup "
          "tiga dimensi komplementer. Pengujian fungsional menggunakan metode "
          "black-box testing terhadap seluruh skenario penggunaan; pengujian "
          "usability menggunakan System Usability Scale (SUS) dengan target "
          "skor minimal 70; serta pengujian performa sistem meliputi waktu "
          "respons, keandalan agen, dan stabilitas integrasi API."),
    ("b", ""),
    ("p", "Tahap kelima adalah Pengembangan Final dan Dokumentasi, mencakup "
          "penyempurnaan sistem berdasarkan hasil evaluasi, penguatan "
          "keamanan kredensial API, serta penyusunan dokumentasi teknis dan "
          "panduan pengguna (User Manual)."),
    ("b", ""),
    ("p", "Tahap keenam mencakup pilot deployment pada lingkungan promosi "
          "digital FSM UNDIP, pelatihan pengguna bagi pengelola kanal "
          "digital fakultas dan departemen, serta pemantauan operasional "
          "awal untuk mengukur efektivitas promosi berbasis data."),
    ("b", ""),
    ("p", "Penelitian ini melibatkan 2 orang peneliti utama dan 12 asisten "
          "peneliti. Pembagian tanggung jawab untuk setiap tahap adalah "
          "sebagai berikut."),
    ("b", ""),
    ("TEAMTABLE", ""),
    ("b", ""),
]

# Tanggung jawab tim (kolom ke-3) untuk 5 baris isi tabel tim
TEAM_RESP = [
    "Arsitektur sistem Agentic AI, orkestrasi OpenClaw, manajemen penelitian "
    "dan koordinasi tim",
    "Desain sistem, integrasi API media sosial, pengujian performa dan "
    "evaluasi sistem",
    "Pengembangan monitoring agent dan content automation agent berbasis "
    "OpenClaw (Python/Node.js)",
    "Integrasi web crawler, API media sosial, basis data, dan dashboard "
    "analitik keterlibatan",
    "Pengujian fungsional, dokumentasi sistem, serta pelatihan pengguna",
]

# Nama kegiatan pada tabel jadwal (7 baris isi: baris r2..r8)
SCHEDULE_KEGIATAN = [
    "Studi literatur, analisis kebutuhan, dan pemetaan kanal digital FSM",
    "Perancangan arsitektur sistem dan orkestrasi Agentic AI (OpenClaw)",
    "Implementasi monitoring agent (web & media sosial) dan integrasi API",
    "Implementasi content automation agent (generasi & penjadwalan konten)",
    "Implementasi analytics agent dan dashboard promosi digital",
    "Pengujian fungsional (black-box), usability (SUS), dan performa",
    "Deployment, pelatihan pengguna, penyusunan laporan, dan publikasi",
]

# --- G. DAFTAR PUSTAKA (gaya Vancouver, 1:1 menggantikan 10 entri sumber) ----
REFERENCES = [
    "Wang L, Ma C, Feng X, Zhang Z, Yang H, Zhang J, et al. A survey on large "
    "language model based autonomous agents. Frontiers of Computer Science. "
    "2024;18(6):186345.",
    "Xi Z, Chen W, Guo X, He W, Ding Y, Hong B, et al. The rise and potential "
    "of large language model based agents: a survey. Science China "
    "Information Sciences. 2025;68(2):121101.",
    "Yao S, Zhao J, Yu D, Du N, Shafran I, Narasimhan K, Cao Y. ReAct: "
    "synergizing reasoning and acting in language models. In: International "
    "Conference on Learning Representations (ICLR); 2023.",
    "Park JS, O'Brien JC, Cai CJ, Morris MR, Liang P, Bernstein MS. "
    "Generative agents: interactive simulacra of human behavior. In: "
    "Proceedings of the 36th Annual ACM Symposium on User Interface Software "
    "and Technology (UIST); 2023. p. 1\u201322.",
    "Brown TB, Mann B, Ryder N, Subbiah M, Kaplan J, Dhariwal P, et al. "
    "Language models are few-shot learners. In: Advances in Neural "
    "Information Processing Systems (NeurIPS); 2020. p. 1877\u20131901.",
    "Kaplan AM, Haenlein M. Users of the world, unite! The challenges and "
    "opportunities of social media. Business Horizons. 2010;53(1):59\u201368. "
    "doi: 10.1016/j.bushor.2009.09.003.",
    "Peruta A, Shields AB. Social media in higher education: understanding "
    "how colleges and universities use Facebook. Journal of Marketing for "
    "Higher Education. 2017;27(1):131\u2013143. "
    "doi: 10.1080/08841241.2016.1212451.",
    "Chaffey D, Ellis-Chadwick F. Digital marketing: strategy, "
    "implementation and practice. 7th ed. Harlow: Pearson Education; 2019.",
    "Todor RD. Marketing automation. Bulletin of the Transilvania University "
    "of Brasov. Economic Sciences. 2016;9(2):87\u201394.",
    "Clifton B. Advanced web metrics with Google Analytics. 3rd ed. "
    "Indianapolis: John Wiley & Sons; 2012.",
]


# =============================================================================
# HELPERS
# =============================================================================

def set_runs_text(p, text):
    """Replace paragraph text using the first run's formatting; clear others."""
    runs = p.runs
    if not runs:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def replace_title_in_pengesahan(p, new_title):
    """In '1.\\tJudul Penelitian\\t :\\t<title>\\t', replace only the title."""
    runs = p.runs
    start = None
    for i, r in enumerate(runs):
        if r.text.strip().startswith("Pengembangan Sistem"):
            start = i
            break
    if start is None:
        return
    end = len(runs)
    for j in range(start + 1, len(runs)):
        if runs[j].text == "\t":
            end = j
            break
    runs[start].text = new_title
    for k in range(start + 1, end):
        runs[k].text = ""


def clear_cell_blocks(cell):
    """Remove all <w:p> and <w:tbl> from a cell, keeping <w:tcPr>."""
    tc = cell._tc
    for child in list(tc.iterchildren()):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            tc.remove(child)


def add_cell_paragraph(cell, text="", bold=False, center=False):
    """Append a Normal-styled paragraph to a cell (inherits TNR12/1.5/justify)."""
    p = cell.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        if bold:
            r.bold = True
    return p


def build_section_cell(cell, items, saved):
    """Rebuild a content cell from a list of (kind, text) items.
    'saved' holds deep-copied drawing/table elements to re-insert."""
    clear_cell_blocks(cell)
    for kind, text in items:
        if kind == "h":
            add_cell_paragraph(cell, text, bold=True)
        elif kind == "p":
            add_cell_paragraph(cell, text)
        elif kind == "b":
            add_cell_paragraph(cell, "")
        elif kind == "CAP":
            add_cell_paragraph(cell, text, center=True)
        elif kind == "CAPB":
            add_cell_paragraph(cell, "", center=True)
        elif kind in ("IMG1", "IMG2"):
            el = saved.get(kind)
            if el is not None:
                cell._tc.append(copy.deepcopy(el))
        elif kind == "TEAMTABLE":
            el = saved.get("TEAMTABLE")
            if el is not None:
                new_el = copy.deepcopy(el)
                cell._tc.append(new_el)
                _update_team_table(Table(new_el, cell))


def _update_team_table(tbl):
    """Update column-3 (Tanggung Jawab) for the 5 data rows; keep names/posisi."""
    # rows: 0 = header, 1..5 = data
    for ridx in range(1, min(6, len(tbl.rows))):
        cell = tbl.rows[ridx].cells[2]
        if ridx - 1 < len(TEAM_RESP):
            # set text on first paragraph, clear rest
            paras = cell.paragraphs
            set_runs_text(paras[0], TEAM_RESP[ridx - 1])
            for extra in paras[1:]:
                set_runs_text(extra, "")


# =============================================================================
# MAIN
# =============================================================================

def main():
    doc = Document(SRC)

    # ---- 1. COVER TITLE -----------------------------------------------------
    for p in doc.paragraphs:
        if "Monitoring Kinerja Pegawai" in p.text and p.alignment is not None:
            set_runs_text(p, TITLE)
            break

    # ---- 2. PENGESAHAN: judul -----------------------------------------------
    for p in doc.paragraphs:
        if p.text.strip().startswith("1.") and "Judul Penelitian" in p.text:
            replace_title_in_pengesahan(p, TITLE)
            break

    t0 = doc.tables[0]

    # ---- 3. ROW 1: JUDUL (di dalam tabel) -----------------------------------
    set_runs_text(t0.rows[1].cells[0].paragraphs[0], TITLE)

    # ---- 4. ROW 3: RINGKASAN ------------------------------------------------
    set_runs_text(t0.rows[3].cells[0].paragraphs[0], RINGKASAN)

    # ---- 5. ROW 5: KATA KUNCI -----------------------------------------------
    set_runs_text(t0.rows[5].cells[0].paragraphs[0], KEYWORDS)

    # ---- 6. ROW 7: PENDAHULUAN (preserve Gambar 1) --------------------------
    cell7 = t0.rows[7].cells[0]
    saved7 = {}
    for child in cell7._tc.iterchildren():
        if child.tag == qn("w:p") and child.findall(".//" + qn("w:drawing")):
            saved7["IMG1"] = copy.deepcopy(child)
            break
    build_section_cell(cell7, PENDAHULUAN, saved7)

    # ---- 7. ROW 9: METODE (preserve Gambar 2 + tabel tim) -------------------
    cell9 = t0.rows[9].cells[0]
    saved9 = {}
    for child in cell9._tc.iterchildren():
        if child.tag == qn("w:p") and child.findall(".//" + qn("w:drawing")):
            saved9["IMG2"] = copy.deepcopy(child)
        elif child.tag == qn("w:tbl"):
            saved9["TEAMTABLE"] = copy.deepcopy(child)
    build_section_cell(cell9, METODE, saved9)

    # ---- 8. ROW 11: JADWAL (update nama kegiatan) ---------------------------
    cell11 = t0.rows[11].cells[0]
    for child in cell11._tc.iterchildren():
        if child.tag == qn("w:tbl"):
            sched = Table(child, cell11)
            # data rows r2..r8 -> kegiatan di kolom 1
            for i, keg in enumerate(SCHEDULE_KEGIATAN):
                ridx = 2 + i
                if ridx < len(sched.rows):
                    c = sched.rows[ridx].cells[1]
                    set_runs_text(c.paragraphs[0], keg)
                    for extra in c.paragraphs[1:]:
                        set_runs_text(extra, "")
            break

    # ---- 9. ROW 14: DAFTAR PUSTAKA (ganti 1:1) ------------------------------
    cell14 = t0.rows[14].cells[0]
    ref_paras = [p for p in cell14.paragraphs]
    n = len(ref_paras)
    for i, ref in enumerate(REFERENCES):
        if i < n:
            set_runs_text(ref_paras[i], ref)
        else:
            add_cell_paragraph(cell14, ref)
    # kosongkan sisa paragraf bila jumlah sumber > referensi baru
    for j in range(len(REFERENCES), n):
        set_runs_text(ref_paras[j], "")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
