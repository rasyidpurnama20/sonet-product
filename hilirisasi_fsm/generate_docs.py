"""
Generator for Proposal Hilirisasi FSM LAPOR.

Produces two files with formal academic layout (mirip File 1 — PROPOSAL Hilirisasi
INFORMATIKA.pdf):
    - proposal-hilirisasi-fsm-lapor.docx
    - proposal-hilirisasi-fsm-lapor.pdf

Run:
    python3 generate_docs.py
"""

from __future__ import annotations

import os
from copy import deepcopy

# =============================================================================
# CONTENT (extracted/condensed from proposal-hilirisasi-fsm-lapor.md)
# =============================================================================

TITLE = (
    "HILIRISASI PRODUK PERANGKAT LUNAK FSM LAPOR: "
    "SISTEM PELAPORAN INSIDEN BERBASIS PROGRESSIVE WEB APP "
    "UNTUK PENGUATAN TATA KELOLA ASET DAN LAYANAN "
    "DI FAKULTAS SAINS DAN MATEMATIKA UNIVERSITAS DIPONEGORO"
)

PENGUSUL_NAME = "Satriawan Rasyid Purnama, S.Kom., M.Cs."
PENGUSUL_NIP = "NIP: 199805212024061001 / NUPTK: 7853776677130152"

PENGESAHAN_ROWS = [
    ("1.", "Judul Penelitian",
     "Hilirisasi Produk Perangkat Lunak FSM LAPOR: Sistem Pelaporan Insiden "
     "Berbasis Progressive Web App untuk Penguatan Tata Kelola Aset dan "
     "Layanan di Fakultas Sains dan Matematika Universitas Diponegoro"),
    ("2.", "Bidang Ilmu", "Informatika / Teknologi Informasi dan Komunikasi"),
    ("3.", "Pengusul", ""),
    ("", "a. Nama Lengkap", "Satriawan Rasyid Purnama, S.Kom., M.Cs."),
    ("", "b. NIP/NUPTK", "199805212024061001 / 7853776677130152"),
    ("", "c. H-Indeks Scopus / ID-ORCID", "6 / 0000-0003-3770-8814"),
    ("", "d. Fakultas/Departemen/Lab", "FSM / Informatika / Komputasi Visual"),
    ("", "e. Pusat Penelitian", "Jl. Prof. Soedarto, S.H. Tembalang, Semarang"),
    ("", "f. Telepon/Faks (kantor)", "(024) 70594104"),
    ("", "g. Telepon/Faks (Rumah)", "-"),
    ("", "h. HP / E-mail",
     "085225257551 / satriawanrasyidp@lecturer.undip.ac.id"),
    ("4.", "Jangka Waktu Kegiatan", "12 Bulan"),
    ("5.", "Lokasi Penelitian",
     "Laboratorium Komputasi Visual, Departemen Informatika, FSM Undip"),
    ("6.", "Biaya yang Diperlukan",
     "Rp 40.000.000,- (empat puluh juta rupiah)"),
    ("7.", "Sumber Dana",
     "Riset Penugasan Hilirisasi FSM Undip Tahun Anggaran 2026"),
]

ABSTRAK_PARAGRAPHS = [
    "Pengelolaan insiden dan kondisi aset di lingkungan fakultas saat ini "
    "masih menghadapi tantangan berupa proses manual yang tidak terstruktur, "
    "ketiadaan sistem pelacakan berbasis bukti foto-geolokasi, serta "
    "keterlambatan penanganan akibat tidak adanya mekanisme Service Level "
    "Agreement (SLA). Permasalahan ini berdampak langsung pada rendahnya "
    "akuntabilitas, sulitnya audit riwayat insiden, dan inefisiensi alokasi "
    "sumber daya pemeliharaan di Fakultas Sains dan Matematika (FSM) "
    "Universitas Diponegoro.",

    "Hilirisasi ini bertujuan membawa prototipe FSM LAPOR—sebuah Progressive "
    "Web App (PWA) berbasis React 18 + Supabase yang mengintegrasikan "
    "geolokasi GPS, kamera perangkat, workflow multi-role (Pelapor, "
    "Pimpinan, Petugas, Superadmin), pelacakan SLA real-time, notifikasi "
    "push, serta modul Survey Kondisi Aset berbasis siklus PDCA—dari level "
    "riset (TKT 5) menuju level siap operasional institusional (TKT 7). "
    "Tahapan hilirisasi mencakup penyempurnaan arsitektur produksi, pilot "
    "deployment di enam departemen FSM (Informatika, Matematika, Statistika, "
    "Fisika, Kimia, Biologi) beserta Tata Usaha Fakultas, pelatihan "
    "operator, pendampingan teknis selama 12 bulan, serta validasi usability "
    "(System Usability Scale ≥ 70) dan performa (Lighthouse Performance ≥ "
    "80). Luaran wajib berupa produk perangkat lunak FSM LAPOR v1.0 yang "
    "deployed dan operasional, satu publikasi jurnal internasional terindeks "
    "Scopus, dan pendaftaran HKI. Proyeksi Revenue Generating Activity (RGA) "
    "jangka menengah berasal dari replikasi sistem ke fakultas/PTN lain "
    "dengan skema sewa-langganan.",
]

KEYWORDS = ("Progressive Web App; Sistem Pelaporan Insiden; Geolokasi GPS; "
            "Manajemen Aset; Multi-Role Workflow; Hilirisasi.")

# Body sections: (heading_level, heading_text, [paragraphs...])
# Level 1 = BAB heading, 2 = subsection (1.1), 3 = sub-sub (a/b/c)
SECTIONS = [
    ("BAB", "BAB I PENDAHULUAN", []),
    ("H2", "1.1 Latar Belakang", [
        "Transformasi digital di sektor pendidikan tinggi telah menjadi "
        "agenda prioritas global. Laporan E-Government Survey 2024 PBB "
        "menegaskan bahwa aksesibilitas layanan publik berbasis digital "
        "merupakan indikator utama kematangan tata kelola modern [1]. Studi "
        "tentang transformasi digital di Indonesia juga menunjukkan bahwa "
        "digitalisasi layanan administrasi meningkatkan efisiensi, "
        "transparansi, dan akuntabilitas tata kelola organisasi secara "
        "signifikan [13]. Namun di tingkat operasional fakultas/departemen, "
        "mekanisme pelaporan insiden atau kerusakan fasilitas umumnya masih "
        "dilakukan secara konvensional melalui laporan lisan, chat tidak "
        "terstruktur, atau formulir fisik.",

        "Ketidakefisienan sistem manual ini berdampak pada lambatnya respons "
        "penanganan, lemahnya akuntabilitas, dan sulitnya pelacakan riwayat "
        "kondisi aset. Survei terhadap sistem manajemen pengaduan "
        "menunjukkan bahwa migrasi ke platform digital meningkatkan "
        "transparansi serta efisiensi interaksi pelapor–pengelola [14]. "
        "Munoz dkk. [3] melalui tinjauan sistematis 76 studi menyimpulkan "
        "bahwa partisipasi pengguna dalam pelaporan masalah infrastruktur "
        "meningkat drastis ketika tersedia platform digital yang mudah "
        "diakses dan memberikan umpan balik status secara real-time.",

        "Progressive Web App (PWA) terbukti menjadi solusi pengembangan "
        "lintas platform yang ekonomis: mendukung akses offline, notifikasi "
        "push, dan instalasi tanpa toko aplikasi. Tinjauan sistematis "
        "Marchetto & Morandini [9] dengan 226 responden mengonfirmasi "
        "kemudahan instalasi dan keterlibatan pengguna PWA yang tinggi. "
        "Biorn-Hansen dkk. [5] menegaskan keunggulan biaya pengembangan, "
        "kemudahan pembaruan, dan jangkauan PWA dibanding aplikasi native. "
        "Malavolta dkk. [4] menunjukkan PWA dengan service worker dapat "
        "menyamai performa native pada konektivitas terbatas—relevan untuk "
        "lingkungan kampus dengan keterbatasan jaringan di area tertentu.",

        "Aspek kritis lain adalah penerapan Service Level Agreement (SLA) "
        "sebagai kontrol waktu penanganan. Swain & Garza [6] membuktikan "
        "faktor kritis pencapaian SLA insiden TI mencakup kejelasan "
        "penugasan, eskalasi otomatis, dan visibilitas status real-time. "
        "Implementasi Role-Based Access Control (RBAC) dalam sistem berbasis "
        "web meningkatkan keamanan data dan efisiensi alur kerja "
        "organisasional [7]. Manajemen kondisi aset menjadi dimensi "
        "tambahan; pendekatan PDCA (Plan-Do-Check-Act) merupakan kerangka "
        "perbaikan berkelanjutan yang diakui untuk pemeliharaan fasilitas "
        "[15], sementara Tuhaise dkk. [8] menegaskan kerangka manajemen "
        "aset berbasis digital twin memungkinkan pemantauan kondisi aset "
        "real-time dan terdokumentasi—hal yang tidak dapat dipenuhi sistem "
        "manual.",

        "Hilirisasi sistem FSM LAPOR menjadi langkah strategis penguatan "
        "infrastruktur digital FSM Undip menuju pengelolaan aset dan "
        "layanan berbasis data. Sistem ini merupakan rangkaian roadmap "
        "riset terapan 2026–2030 di bidang smart campus services. Melalui "
        "pilot di enam departemen FSM dan unit Tata Usaha Fakultas, sistem "
        "akan diuji performa teknis, usability, dan validitas workflow-nya. "
        "Dengan pendampingan teknis 12 bulan, hasil pengembangan "
        "ditargetkan mencapai TKT 7 dan siap diadopsi secara institusional, "
        "sekaligus menjadi model implementasi yang dapat direplikasi ke "
        "fakultas lain di Undip dan PTN sejawat.",
    ]),
    ("H2", "1.2 Hasil Riset Awal (Sumber Hilirisasi)", [
        "Riset awal yang menjadi dasar hilirisasi ini adalah pengembangan "
        "prototipe FSM LAPOR v0.x oleh tim peneliti Departemen Informatika "
        "FSM Undip pada 2025–2026 dengan judul 'FSM LAPOR: Pengembangan "
        "Sistem Pelaporan Insiden Berbasis Progressive Web App pada "
        "Fakultas Sains dan Matematika Universitas Diponegoro'. Prototipe "
        "tersebut menghasilkan: (1) Arsitektur PWA berbasis Vite 5 + React "
        "18 + TypeScript di sisi frontend dan Supabase (PostgreSQL + Auth + "
        "Storage + Realtime) di sisi backend sebagai platform BaaS "
        "open-source; (2) Skema basis data relasional dengan Row Level "
        "Security (RLS) untuk enkapsulasi hak akses per-role, "
        "mengimplementasikan prinsip RBAC [7]; (3) Modul inti pelaporan "
        "insiden berbasis foto + geolokasi GPS, workflow lima-state "
        "(Dikirim → Diterima → Ditugaskan → Diselesaikan → Diverifikasi), "
        "SLA countdown dan eskalasi otomatis [6], notifikasi real-time via "
        "Supabase Realtime Channels berbasis WebSocket, serta modul Survey "
        "Kondisi Aset siklus PDCA [15]; (4) Validasi awal pada lingkup "
        "Departemen Informatika menunjukkan kelayakan teknis sebagai PWA "
        "installable dan offline-capable.",

        "Prototipe ini memiliki kebaruan berupa integrasi tunggal: "
        "pelaporan foto-GPS, workflow multi-role + SLA tracking, "
        "notifikasi push real-time, dan modul survey aset PDCA dalam satu "
        "PWA institusional [3], [9], [12]; serta arsitektur "
        "BaaS-Supabase + RLS PostgreSQL untuk pelaporan institusional yang "
        "belum terdokumentasi luas secara akademis. Riset terkait tim "
        "peneliti yang turut menjadi penguat keyakinan hilirisasi mencakup "
        "studi pelaporan keluhan berbasis deep learning [12] dan tinjauan "
        "PWA [4], [5], [9], [19].",

        "Hilirisasi pada proposal ini diarahkan untuk: (a) memperluas "
        "cakupan dari satu departemen menjadi seluruh FSM (6 departemen + "
        "Tata Usaha Fakultas), (b) penguatan keamanan & keandalan "
        "produksi (RLS lanjutan, audit trail, backup otomatis), (c) "
        "integrasi dashboard analitik lintas departemen, (d) pelatihan "
        "operator, dan (e) pendampingan operasional 12 bulan untuk "
        "mencapai TKT 7.",
    ]),
    ("H2", "1.3 Tujuan Penelitian", [
        "Tujuan hilirisasi ini meliputi:",
    ]),
    ("OL", "tujuan", [
        "Mengembangkan dan menyempurnakan produk perangkat lunak FSM LAPOR "
        "sebagai PWA siap-produksi yang mendukung pelaporan insiden "
        "berbasis foto + geolokasi GPS secara real-time.",
        "Mendesain dan mengimplementasikan arsitektur sistem yang terdiri "
        "dari PWA Frontend Engine, Supabase Backend & API Engine, dan "
        "Dashboard Analytics Engine untuk mendukung integrasi data, "
        "keamanan (RLS), dan kemudahan akses bagi tiap departemen di FSM.",
        "Melakukan uji validasi teknis terhadap performa sistem meliputi "
        "response time, kestabilan realtime channel, akurasi GPS/kamera, "
        "serta keandalan SLA tracking dan eskalasi otomatis.",
        "Mengimplementasikan sistem secara bertahap pada 6 departemen FSM "
        "(Informatika, Matematika, Statistika, Fisika, Kimia, Biologi) dan "
        "Tata Usaha Fakultas FSM, termasuk pelatihan operator dan "
        "pendampingan teknis.",
        "Menyusun dokumentasi teknis, panduan pengguna (user manual), dan "
        "standar operasional prosedur (SOP) untuk mendukung keberlanjutan "
        "dan skalabilitas implementasi.",
        "Mencapai Tingkat Kesiapan Teknologi (TKT) 7 sebagai indikator "
        "kesiapan produk untuk digunakan secara institusional di FSM Undip.",
        "Menyusun model hilirisasi berkelanjutan yang dapat direplikasi ke "
        "fakultas lain di Undip maupun perguruan tinggi lain di Indonesia.",
    ]),
    ("H2", "1.4 Manfaat Penelitian", [
        "Manfaat penelitian ini meliputi:",
    ]),
    ("UL", "manfaat", [
        "Memberikan kontribusi ilmiah dalam penerapan PWA, RBAC berbasis "
        "RLS, dan siklus PDCA untuk tata kelola aset institusi pendidikan "
        "tinggi.",
        "Menghasilkan sistem perangkat lunak inovatif yang mempersingkat "
        "siklus pelaporan–penanganan insiden dan meningkatkan akuntabilitas "
        "pemeliharaan fasilitas FSM.",
        "Mendukung kebijakan data-driven faculty melalui dashboard yang "
        "menampilkan tren insiden, SLA compliance, dan kondisi aset lintas "
        "departemen secara real-time.",
        "Menjadi produk hilirisasi riset terapan informatika yang dapat "
        "diadopsi fakultas lain sebagai model sistem pelaporan "
        "institusional.",
        "Meningkatkan kapasitas analisis data dan kesiapan digital unit "
        "kerja (Tata Usaha, Sub-bagian Umum & BMN, Pimpinan Departemen) "
        "untuk merespons isu pemeliharaan secara cepat dan berbasis bukti.",
        "Memperkuat infrastruktur layanan FSM melalui sistem pelaporan "
        "terpusat dengan mekanisme eskalasi otomatis.",
        "Memberikan manfaat sosial-kelembagaan: meningkatkan transparansi, "
        "akuntabilitas, dan kepercayaan civitas akademika terhadap layanan "
        "fakultas.",
        "Mendorong ekosistem inovasi berkelanjutan di bidang informatika "
        "yang menghubungkan riset akademik dengan kebutuhan riil institusi.",
    ]),

    ("BAB", "BAB II MITRA DAN TARGET PENGGUNAAN", []),
    ("H2", "2.1 Mitra untuk Hilirisasi", [
        "Hilirisasi FSM LAPOR dilaksanakan melalui kemitraan strategis "
        "pada dua sisi:",
        "Sisi penyedia data & infrastruktur teknis: Supabase (PostgreSQL + "
        "Auth + Storage + Realtime) sebagai Backend-as-a-Service utama; "
        "Vercel/Cloudflare Pages sebagai hosting PWA frontend dengan "
        "dukungan edge network; layanan notifikasi web push (mis. Firebase "
        "Cloud Messaging atau OneSignal) untuk push notification lintas "
        "perangkat; PT/registrar domain .id untuk penyediaan domain "
        "institusional dan SSL certificate.",
        "Sisi pengguna institusional (mitra utama hilirisasi di FSM Undip): "
        "Dekanat FSM Undip (cq. Wakil Dekan Bidang II/Sumber Daya) sebagai "
        "executive sponsor; 6 Departemen di FSM (Informatika, Matematika, "
        "Statistika, Fisika, Kimia, Biologi); Tata Usaha Fakultas FSM dan "
        "Sub-bagian Umum & BMN sebagai pemilik workflow pemeliharaan, "
        "aset, dan logistik; Unit Pemeliharaan/Layanan Teknis FSM sebagai "
        "end user role Petugas; Koordinator Laboratorium di tiap "
        "departemen sebagai power user untuk Survey Kondisi Aset.",
        "Dengan model multipihak ini, hilirisasi diharapkan memperkuat "
        "ekosistem riset terapan FSM melalui sinergi akademisi, penyedia "
        "teknologi, dan pengambil kebijakan di tingkat fakultas.",
    ]),
    ("H2", "2.2 Target Penggunaan", [
        "Produk FSM LAPOR ditargetkan menjadi infrastruktur strategis FSM "
        "Undip untuk pengelolaan insiden, layanan, dan kondisi aset "
        "berbasis data. Sistem dirancang untuk empat role utama "
        "sebagaimana ditampilkan pada Tabel 2.1.",
    ]),
    ("TABLE_ROLE", "Tabel 2.1 Role Pengguna FSM LAPOR", []),
    ("P", "", [
        "Target penggunaan jangka pendek adalah enam departemen FSM + Tata "
        "Usaha Fakultas. Target jangka menengah adalah replikasi ke "
        "fakultas lain di Undip, sementara jangka panjang adalah model "
        "rujukan nasional sistem pelaporan insiden + manajemen aset "
        "berbasis PWA untuk PTN/PTS lain.",
    ]),
    ("H2", "2.3 Dampak dan Manfaat Hilirisasi", [
        "Dampak dan manfaat hilirisasi mencakup beberapa dimensi berikut:",
    ]),
    ("UL", "dampak", [
        "Kelembagaan: FSM memiliki infrastruktur digital terintegrasi "
        "untuk pelaporan insiden, eskalasi otomatis berbasis SLA, dan "
        "pemantauan kondisi aset yang terdokumentasi siklus PDCA.",
        "Operasional: Pimpinan dan Tata Usaha memperoleh dashboard "
        "real-time (jumlah laporan, SLA compliance rate, kategori insiden "
        "dominan, kondisi aset per departemen) untuk pengambilan keputusan "
        "berbasis data; tim Petugas memperoleh sistem penugasan yang "
        "transparan dengan audit trail.",
        "Akademik: Memperkuat kapasitas riset terapan Departemen "
        "Informatika dalam pengembangan PWA, BaaS-architecture, dan "
        "workflow engineering untuk konteks lokal Indonesia; menghasilkan "
        "publikasi jurnal internasional terindeks Scopus.",
        "Civitas Akademika: Meningkatkan kepuasan layanan, transparansi "
        "penanganan keluhan, dan rasa kepemilikan terhadap fasilitas "
        "fakultas.",
        "Jangka Panjang: Menjadi model nasional sistem smart-campus "
        "reporting berbasis PWA, mendukung kolaborasi lintas-PTN, dan "
        "memperkuat peran FSM Undip sebagai pionir hilirisasi informatika "
        "untuk tata kelola fakultas modern.",
    ]),

    ("BAB", "BAB III METODOLOGI HILIRISASI", []),
    ("H2", "3.1 Tahapan Kegiatan", [
        "Hilirisasi dilakukan melalui lima tahapan utama:",
    ]),
    ("OL", "tahapan", [
        "Penyempurnaan Desain Sistem dan Arsitektur Teknis. Evaluasi "
        "prototipe v0.x dan penyempurnaan arsitektur agar memenuhi standar "
        "produksi: hardening RLS PostgreSQL, audit trail lengkap, backup "
        "terjadwal, observability (log + metrik), dan service worker PWA "
        "untuk offline capability + auto-update. Arsitektur final "
        "mencakup tiga komponen: PWA Frontend Engine (Vite 5 + React 18 + "
        "TypeScript), Supabase Backend & API Engine (PostgreSQL + Auth + "
        "Storage + Realtime + RLS), dan Dashboard Analytics Engine "
        "(visualisasi tren insiden, SLA, dan kondisi aset).",
        "Pengembangan dan Integrasi Skala Fakultas. Pengembangan modul "
        "level produksi dalam tiga sprint iteratif: Sprint 1 (autentikasi "
        "multi-role, manajemen laporan inti, integrasi Geolocation API, "
        "akses kamera perangkat), Sprint 2 (workflow lima-state + "
        "multi-assignee, SLA tracking + eskalasi otomatis, notifikasi "
        "real-time via Supabase Realtime Channels + web push), Sprint 3 "
        "(modul Survey Kondisi Aset PDCA, dashboard statistik & analitik "
        "lintas departemen, master data management, konfigurasi PWA — "
        "service worker, manifest, auto-update).",
        "Uji Coba dan Validasi Fungsional (Pilot Test). Pilot deployment "
        "di 6 departemen FSM + Tata Usaha Fakultas. Kegiatan: pelatihan "
        "operator, uji dashboard, pengujian kecepatan pemrosesan, validasi "
        "kesesuaian workflow dengan praktik nyata.",
        "Validasi Model Bisnis dan Mekanisme Pemanfaatan. Penyusunan "
        "model operasional di FSM: skema pemeliharaan, kontrol akses, "
        "mekanisme pendanaan tahunan, serta perhitungan proyeksi RGA "
        "untuk replikasi ke fakultas/PTN lain.",
        "Perizinan dan Dokumentasi Teknis. Dokumentasi teknis (SRS sesuai "
        "IEEE 830, ADR, deployment guide), laporan hasil uji coba, "
        "penyelarasan kebijakan keamanan & privasi (PDP) sesuai pedoman "
        "Undip, serta pengajuan HKI.",
    ]),
    ("H2", "3.2 Metode Pengujian dan Validasi", [
        "Metode pengujian dilakukan untuk memastikan kelayakan teknis, "
        "fungsional, dan operasional sistem. Pengujian dilakukan melalui "
        "beberapa pendekatan berikut:",
    ]),
    ("OL", "uji", [
        "Uji Teknis dan Performa Sistem. Pengukuran response time API & "
        "realtime channel, stabilitas server Supabase, efisiensi bundle "
        "PWA, dan offline capability via simulasi network throttling. "
        "Skor Google Lighthouse Performance ≥ 80 ditetapkan sebagai "
        "ambang.",
        "Uji Validasi Akurasi Fungsional. Black-box testing terhadap "
        "seluruh use-case SRS (alur lima-state untuk empat role), "
        "validasi akurasi geolokasi GPS, dan validasi mekanisme eskalasi "
        "SLA dengan dataset skenario sintetis.",
        "Uji Coba Lapangan (Pilot Implementation). Implementasi di 6 "
        "departemen FSM + Tata Usaha Fakultas. Usability diuji "
        "menggunakan System Usability Scale (SUS) [17], [18] dengan "
        "minimal 25 responden mewakili keempat role; ambang SUS ≥ 70 "
        "(grade C/acceptable).",
        "Uji Keamanan dan Keandalan Sistem. Penilaian autentikasi "
        "(Supabase Auth), kontrol akses berbasis RLS PostgreSQL [7], "
        "proteksi data sensitif (foto bukti, lokasi GPS), serta "
        "pengujian backup-restore.",
    ]),
    ("H2", "3.3 Indikator Kinerja dan Ketercapaian", [
        "Keberhasilan hilirisasi diukur melalui capaian teknis, "
        "fungsional, dan kelembagaan sebagaimana terangkum pada "
        "Tabel 3.1.",
    ]),
    ("TABLE_KPI", "Tabel 3.1 Indikator Kinerja dan Target Capaian", []),

    ("BAB", "BAB IV RENCANA LUARAN", []),
    ("H2", "4.1 Luaran", [
        "Produk yang dikembangkan adalah perangkat lunak FSM LAPOR — "
        "Sistem Pelaporan Insiden & Manajemen Aset berbasis Progressive "
        "Web App — yang berfungsi sebagai infrastruktur digital fakultas "
        "untuk pengelolaan insiden, eskalasi berbasis SLA, dan "
        "pemantauan kondisi aset (Tabel 4.1).",
    ]),
    ("TABLE_PRODUK", "Tabel 4.1 Jenis Produk dan Fungsi Utama", []),
    ("P", "", [
        "Daftar luaran yang ditargetkan dijabarkan pada Tabel 4.2.",
    ]),
    ("TABLE_LUARAN", "Tabel 4.2 Luaran Wajib dan Luaran Tambahan", []),

    ("BAB", "BAB V RENCANA ANGGARAN DAN JADWAL PELAKSANAAN", []),
    ("H2", "5.1 Rencana Anggaran Biaya (RAB)", [
        "Total pagu yang diusulkan adalah Rp 40.000.000,- (empat puluh "
        "juta rupiah), dialokasikan secara seimbang antara belanja "
        "personil dan belanja operasional.",
    ]),
    ("H3", "A. BELANJA PERSONIL (Honorarium Asisten/Mahasiswa Pembantu)", []),
    ("TABLE_RAB_A", "", []),
    ("H3", "B. BELANJA OPERASIONAL", []),
    ("TABLE_RAB_B", "", []),
    ("H3", "REKAPITULASI", []),
    ("TABLE_RAB_REKAP", "", []),
    ("H2", "5.2 Jadwal Pelaksanaan (Time Schedule)", [
        "Jadwal pelaksanaan kegiatan hilirisasi selama 12 bulan disajikan "
        "pada Tabel 5.1. Tanda ■ menandakan bulan aktif kegiatan.",
    ]),
    ("TABLE_SCHEDULE", "Tabel 5.1 Jadwal Pelaksanaan Hilirisasi", []),
]

# Tables
ROLE_TABLE = [
    ["Role", "Pengguna Utama", "Fungsi Utama"],
    ["Pelapor", "Dosen, mahasiswa, tendik FSM",
     "Membuat laporan insiden dengan foto + GPS, memantau status, "
     "memberi verifikasi penyelesaian."],
    ["Pimpinan", "Kadep, Sekdep, KTU, Pimpinan Fakultas",
     "Menerima laporan, men-disposisi ke Petugas, memantau dashboard "
     "dan SLA compliance."],
    ["Petugas", "Unit Pemeliharaan, teknisi laboratorium",
     "Mengeksekusi penanganan, mengunggah bukti penyelesaian, "
     "memperbarui status."],
    ["Superadmin", "Tim TI Fakultas / Tim Hilirisasi",
     "Mengelola master data, role, kategori insiden, target SLA, "
     "audit trail."],
]

KPI_TABLE = [
    ["No", "Indikator", "Target"],
    ["1", "Tingkat Kesiapan Teknologi (TKT)", "Naik dari TKT 5 → TKT 7"],
    ["2", "Cakupan implementasi",
     "Minimal 6 departemen FSM + Tata Usaha Fakultas"],
    ["3", "Usability (SUS)", "≥ 70 (acceptable)"],
    ["4", "Performa PWA (Lighthouse)",
     "Performance ≥ 80; Accessibility ≥ 90; Best Practices ≥ 90"],
    ["5", "Response time API",
     "< 1 detik (p95) untuk operasi CRUD utama"],
    ["6", "SLA compliance rate pasca-implementasi",
     "≥ 80% laporan tertangani sesuai SLA"],
    ["7", "Dokumentasi & SOP",
     "SRS (IEEE 830), user manual per-role, SOP operasional, "
     "deployment guide"],
    ["8", "Luaran wajib",
     "1 publikasi jurnal internasional terindeks Scopus, 1 produk "
     "PWA deployed, 1 HKI"],
]

PRODUK_TABLE = [
    ["No", "Jenis Produk", "Fungsi Utama"],
    ["1", "PWA Pelaporan Insiden FSM LAPOR",
     "Memungkinkan civitas akademika FSM membuat laporan insiden "
     "berbasis foto + geolokasi GPS, memantau status secara real-time, "
     "dan menerima notifikasi push untuk perubahan status."],
    ["2", "Dashboard Manajemen & Analitik FSM LAPOR",
     "Menggabungkan data laporan, SLA, dan survey aset PDCA dari "
     "seluruh departemen FSM ke dalam satu dashboard interaktif: tren "
     "insiden, SLA compliance, kondisi aset per-departemen, serta "
     "audit trail lengkap untuk pengambilan keputusan strategis "
     "pimpinan fakultas."],
]

LUARAN_TABLE = [
    ["Kategori", "Luaran"],
    ["Luaran Wajib",
     "(1) Produk perangkat lunak FSM LAPOR v1.0 deployed dan "
     "operasional di 6 departemen FSM + Tata Usaha Fakultas; "
     "(2) Laporan teknis hilirisasi (arsitektur, hasil validasi, "
     "deployment guide); (3) Laporan penerapan & uji coba lapangan; "
     "(4) Panduan pengguna (user manual) per-role + SOP institusional; "
     "(5) Laporan evaluasi performa & umpan balik pengguna "
     "(SUS + Lighthouse)."],
    ["Luaran Tambahan",
     "(1) Publikasi jurnal internasional terindeks Scopus (target: "
     "jurnal Q3/Q2 di bidang software engineering / information "
     "systems); (2) Pendaftaran HKI untuk perangkat lunak FSM LAPOR; "
     "(3) Proyeksi RGA dari skema replikasi/sewa-langganan: estimasi "
     "Rp 5.000.000/fakultas/tahun × 10 fakultas = Rp 50.000.000/tahun "
     "sebagai potensi pendapatan jangka menengah pasca-hilirisasi."],
]

RAB_A_TABLE = [
    ["No", "Uraian", "Volume", "Satuan", "Harga Satuan (Rp)", "Jumlah (Rp)"],
    ["1", "Honorarium asisten Frontend Developer (PWA React/TypeScript)",
     "240", "OJ", "25.000", "6.000.000"],
    ["2", "Honorarium asisten Backend Developer (Supabase/PostgreSQL/RLS)",
     "240", "OJ", "25.000", "6.000.000"],
    ["3", "Honorarium asisten QA & UI/UX Tester",
     "160", "OJ", "25.000", "4.000.000"],
    ["4", "Honorarium asisten Helpdesk, Pelatihan & Sosialisasi",
     "160", "OJ", "25.000", "4.000.000"],
    ["", "Sub Total A", "", "", "", "20.000.000"],
]

RAB_B_TABLE = [
    ["No", "Uraian", "Volume", "Satuan", "Harga Satuan (Rp)", "Jumlah (Rp)"],
    ["1", "Supabase Pro (Auth + DB + Storage + Realtime) — backend produksi",
     "12", "bulan", "600.000", "7.200.000"],
    ["2",
     "Hosting PWA frontend (Vercel/Cloudflare Pages tier produksi)",
     "12", "bulan", "250.000", "3.000.000"],
    ["3", "Domain .id + SSL Certificate", "1", "tahun", "500.000",
     "500.000"],
    ["4", "Layanan Web Push Notification (FCM/OneSignal)",
     "12", "bulan", "200.000", "2.400.000"],
    ["5", "UI library/template & icon set berlisensi", "1", "paket",
     "1.500.000", "1.500.000"],
    ["6",
     "Workshop & pelatihan operator (6 departemen + Tata Usaha)",
     "7", "paket", "350.000", "2.450.000"],
    ["7", "ATK, print & jilid laporan/SOP/user manual", "1", "paket",
     "950.000", "950.000"],
    ["8", "Transportasi & koordinasi lapangan (pilot 6 departemen)",
     "1", "paket", "1.000.000", "1.000.000"],
    ["9",
     "Biaya publikasi jurnal internasional terindeks Scopus (APC)",
     "1", "kali", "1.000.000", "1.000.000"],
    ["", "Sub Total B", "", "", "", "20.000.000"],
]

RAB_REKAP_TABLE = [
    ["No", "Komponen", "Jumlah (Rp)", "Persentase"],
    ["A", "Belanja Personil (Honorarium)", "20.000.000", "50%"],
    ["B", "Belanja Operasional", "20.000.000", "50%"],
    ["", "TOTAL KESELURUHAN", "40.000.000", "100%"],
]

SCHEDULE_TABLE_HEADER = (
    ["No", "Kegiatan"] + [f"B{i}" for i in range(1, 13)]
)
SCHEDULE_TABLE_ROWS = [
    ["1",
     "Penyempurnaan desain sistem & arsitektur teknis "
     "(hardening RLS, audit trail, service worker)",
     "■", "■", "", "", "", "", "", "", "", "", "", ""],
    ["2",
     "Pengembangan & integrasi skala fakultas "
     "(Sprint 1–3: PWA, workflow, SLA, push, dashboard, survey PDCA)",
     "", "■", "■", "■", "■", "", "", "", "", "", "", ""],
    ["3",
     "Uji coba & validasi fungsional (pilot di 6 departemen FSM "
     "+ Tata Usaha Fakultas)",
     "", "", "", "", "■", "■", "■", "", "", "", "", ""],
    ["4",
     "Pendampingan operasional & evaluasi performa "
     "(SUS, Lighthouse, SLA compliance)",
     "", "", "", "", "", "■", "■", "■", "■", "■", "■", "■"],
    ["5",
     "Dokumentasi teknis (SRS IEEE 830, user manual, SOP, "
     "deployment guide)",
     "", "", "", "■", "■", "■", "■", "", "", "", "", ""],
    ["6",
     "Penyusunan laporan akhir, publikasi jurnal Scopus, "
     "pengajuan HKI, & proyeksi RGA",
     "", "", "", "", "", "", "", "■", "■", "■", "■", "■"],
]

REFERENCES = [
    "[1] United Nations, E-Government Survey 2024: Accelerating Digital "
    "Transformation for Sustainable Development, United Nations Department "
    "of Economic and Social Affairs, New York, 2024.",
    "[2] H. Phour, D. Sharma, and N. S. Talwandi, \"Crowdsourcing "
    "Applications in Smart Cities,\" in Intelligent Systems Design and "
    "Applications. ISDA 2023, Lecture Notes in Networks and Systems, "
    "vol. 1049, A. Abraham et al., Eds. Cham: Springer, 2024, pp. 25–36. "
    "doi: 10.1007/978-3-031-64779-6_3.",
    "[3] P. Munoz, S. Casademont, and F. Marques, \"Smart City Applications "
    "to Promote Citizen Participation in City Management and Governance: A "
    "Systematic Review,\" Informatics, vol. 9, no. 4, p. 89, Oct. 2022. "
    "doi: 10.3390/informatics9040089.",
    "[4] I. Malavolta, G. Procaccianti, P. Noorland, and P. Vukmirovic, "
    "\"Assessing the Impact of Service Workers on the Energy Efficiency of "
    "Progressive Web Apps,\" in Proc. 2017 IEEE/ACM 4th Int. Conf. Mobile "
    "Software Engineering and Systems (MOBILESoft), IEEE, 2017, pp. 35–45.",
    "[5] A. Biorn-Hansen, T. A. Majchrzak, and T.-M. Gronli, \"Progressive "
    "Web Apps: The Possible Web-native Unifier for Mobile Development,\" "
    "in Proc. 13th Int. Conf. Web Information Systems and Technologies "
    "(WEBIST), 2017, pp. 344–351. doi: 10.5220/0006728803440351.",
    "[6] A. K. Swain and V. R. Garza, \"Key Factors in Achieving Service "
    "Level Agreements (SLA) for Information Technology (IT) Incident "
    "Resolution,\" Information Systems Frontiers, vol. 25, no. 2, "
    "pp. 819–834, 2023. doi: 10.1007/s10796-022-10266-5.",
    "[7] Z. M. Iqal, A. Selamat, and O. Krejcar, \"A Comprehensive "
    "Systematic Review of Access Control in IoT: Requirements, "
    "Technologies, and Evaluation Metrics,\" IEEE Access, vol. 12, "
    "pp. 12636–12654, 2024. doi: 10.1109/ACCESS.2023.3347495.",
    "[8] V. V. Tuhaise, J. H. M. Tah, and F. H. Abanda, \"Technologies "
    "for Digital Twin Applications in Construction,\" Automation in "
    "Construction, vol. 152, p. 104931, Aug. 2023. "
    "doi: 10.1016/j.autcon.2023.104931.",
    "[9] T. Marchetto and M. Morandini, \"User Perceptions of Progressive "
    "Web App Features: An Analytical Approach and a Systematic Literature "
    "Review,\" in Proc. 9th Int. Congr. Information and Communication "
    "Technology (ICICT 2024), Lecture Notes in Networks and Systems, "
    "vol. 1001. Singapore: Springer, 2024, pp. 163–172. "
    "doi: 10.1007/978-981-97-4581-4_14.",
    "[10] M. Pan, \"Prototyping Methods: Techniques and its "
    "Significance,\" Journal of Research and Development, vol. 11, "
    "p. 216, Jun. 2023.",
    "[11] M. Pan, \"Prototyping Methods: Techniques and its "
    "Significance,\" J. Res. Dev., vol. 11, p. 216, Jun. 2023. "
    "doi: 10.35248/2311-3278.23.11.216.",
    "[12] F. Shama, A. Aziz, and L. B. M. Deya, \"CitySolution: A "
    "Complaining Task Distributive Mobile Application for Smart City "
    "Corporation Using Deep Learning,\" arXiv preprint arXiv:2410.12882, "
    "2024.",
    "[13] M. Sebő and G. Bel, \"E-Government and Provision of Public "
    "Services: Economic, Social, and Political Determinants of Citizen "
    "Complaints,\" International Public Management Journal, vol. 27, "
    "no. 4, pp. 659–679, 2024. doi: 10.1080/10967494.2023.2273343.",
    "[14] D. A. Puspitasari and T. Kurniawan, \"Assessing the National "
    "Complaint Handling System in Indonesia (LAPOR!) Using the "
    "Design-Reality Gap Model,\" International Journal of Electronic "
    "Governance, vol. 15, no. 2, pp. 118–134, 2023. "
    "doi: 10.1504/IJEG.2023.132329.",
    "[15] E. Üstündağlı Erten, \"Complaint Management through the E-State "
    "Portal: Is Digitalization Actually Beneficial?\" Proceedings, "
    "vol. 101, no. 1, p. 1, 2024. doi: 10.3390/proceedings2024101001.",
    "[16] D. Kumar et al., \"Digital Twins in the Construction Industry: "
    "A Comprehensive Review of Current Implementations, Enabling "
    "Technologies, and Future Directions,\" Sustainability, vol. 15, "
    "no. 14, p. 10908, 2023. doi: 10.3390/su151410908.",
    "[17] P. Vlachogianni and N. Tselios, \"Perceived Usability "
    "Evaluation of Educational Technology Using the System Usability "
    "Scale (SUS): A Systematic Review,\" Journal of Research on "
    "Technology in Education, vol. 54, no. 3, pp. 394–410, 2022. "
    "doi: 10.1080/15391523.2020.1867938.",
    "[18] O. Suria, \"A Statistical Analysis of System Usability Scale "
    "(SUS) Evaluations in Online Learning Platform,\" Journal of "
    "Information Systems and Informatics, vol. 6, no. 2, pp. 992–1007, "
    "2024. doi: 10.51519/journalisi.v6i2.750.",
    "[19] S. Huber, L. Demetz, and M. Felderer, \"A Comparative Study on "
    "the Energy Consumption of Progressive Web Apps,\" Information "
    "Systems, vol. 108, p. 102017, Sep. 2022. "
    "doi: 10.1016/j.is.2022.102017.",
]


TABLE_REGISTRY = {
    "TABLE_ROLE": ROLE_TABLE,
    "TABLE_KPI": KPI_TABLE,
    "TABLE_PRODUK": PRODUK_TABLE,
    "TABLE_LUARAN": LUARAN_TABLE,
    "TABLE_RAB_A": RAB_A_TABLE,
    "TABLE_RAB_B": RAB_B_TABLE,
    "TABLE_RAB_REKAP": RAB_REKAP_TABLE,
    "TABLE_SCHEDULE": [SCHEDULE_TABLE_HEADER] + SCHEDULE_TABLE_ROWS,
}

# =============================================================================
# DOCX BUILDER
# =============================================================================

def build_docx(out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins similar to academic standard (4-3-3-3 cm)
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    def set_paragraph_format(p, line_spacing=1.5, before=0, after=6,
                             align=None, first_line_indent=None):
        pf = p.paragraph_format
        pf.line_spacing = line_spacing
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        if align is not None:
            p.alignment = align
        if first_line_indent is not None:
            pf.first_line_indent = Cm(first_line_indent)

    def add_para(text, *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 line_spacing=1.5, before=0, after=6, indent=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        set_paragraph_format(p, line_spacing=line_spacing, before=before,
                             after=after, align=align,
                             first_line_indent=indent)
        return p

    def set_cell_borders(cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "000000")
            tc_borders.append(border)
        tc_pr.append(tc_borders)

    def shade_cell(cell, color="D9D9D9"):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tc_pr.append(shd)

    def add_bordered_table(rows, header=True, col_widths=None, font_size=11,
                           header_bold=True):
        n_cols = len(rows[0])
        tbl = doc.add_table(rows=len(rows), cols=n_cols)
        tbl.autofit = False
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = tbl.cell(i, j)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_borders(cell)
                if header and i == 0:
                    shade_cell(cell, "D9D9D9")
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = "Times New Roman"
                run.font.size = Pt(font_size)
                if header and i == 0 and header_bold:
                    run.bold = True
                pf = p.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(2)
                pf.line_spacing = 1.15
                if header and i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if col_widths and j < len(col_widths):
                    cell.width = col_widths[j]
        return tbl

    # ----- COVER PAGE --------------------------------------------------------
    add_para("PROPOSAL", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=0)
    add_para("RISET PENUGASAN HILIRISASI", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=18)
    for _ in range(2):
        add_para("", after=0)
    add_para(TITLE, bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=24)
    for _ in range(2):
        add_para("", after=0)
    add_para("Pengusul:", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=6)
    add_para(PENGUSUL_NAME, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=0)
    add_para(PENGUSUL_NIP, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=24)
    for _ in range(6):
        add_para("", after=0)
    add_para("FAKULTAS SAINS DAN MATEMATIKA", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=0)
    add_para("UNIVERSITAS DIPONEGORO", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=0)
    add_para("TAHUN 2026", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, after=0)
    doc.add_page_break()

    # ----- HALAMAN PENGESAHAN ------------------------------------------------
    add_para("HALAMAN PENGESAHAN", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    pengesahan_tbl = doc.add_table(rows=len(PENGESAHAN_ROWS), cols=3)
    pengesahan_tbl.autofit = False
    for i, (no, label, value) in enumerate(PENGESAHAN_ROWS):
        for j, val in enumerate([no, label, value]):
            cell = pengesahan_tbl.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            pf = p.paragraph_format
            pf.line_spacing = 1.15
            pf.space_after = Pt(2)
        pengesahan_tbl.cell(i, 0).width = Cm(1)
        pengesahan_tbl.cell(i, 1).width = Cm(5.5)
        pengesahan_tbl.cell(i, 2).width = Cm(8)

    add_para("", after=12)
    add_para("Semarang, ........... 2026", size=12,
             align=WD_ALIGN_PARAGRAPH.RIGHT, after=12)

    sig_tbl = doc.add_table(rows=4, cols=2)
    sig_data = [
        ("Menyetujui,", "Ketua Pengusul,"),
        ("Dekan FSM UNDIP", ""),
        ("", ""),
        ("Prof. Dr. Kusworo Adi, S.Si., M.T.",
         "Satriawan Rasyid Purnama, S.Kom., M.Cs."),
    ]
    nip_data = ("NIP. 197203171998021001", "NIP. 199805212024061001")
    for i, (a, b) in enumerate(sig_data):
        for j, val in enumerate([a, b]):
            cell = sig_tbl.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            if i == 3:
                run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_tbl_row = sig_tbl.add_row()
    for j, val in enumerate(nip_data):
        cell = sig_tbl_row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ----- DAFTAR ISI (placeholder simple) -----------------------------------
    add_para("DAFTAR ISI", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    toc_entries = [
        "ABSTRAK",
        "BAB I PENDAHULUAN",
        "    1.1 Latar Belakang",
        "    1.2 Hasil Riset Awal (Sumber Hilirisasi)",
        "    1.3 Tujuan Penelitian",
        "    1.4 Manfaat Penelitian",
        "BAB II MITRA DAN TARGET PENGGUNAAN",
        "    2.1 Mitra untuk Hilirisasi",
        "    2.2 Target Penggunaan",
        "    2.3 Dampak dan Manfaat Hilirisasi",
        "BAB III METODOLOGI HILIRISASI",
        "    3.1 Tahapan Kegiatan",
        "    3.2 Metode Pengujian dan Validasi",
        "    3.3 Indikator Kinerja dan Ketercapaian",
        "BAB IV RENCANA LUARAN",
        "    4.1 Luaran",
        "BAB V RENCANA ANGGARAN DAN JADWAL PELAKSANAAN",
        "    5.1 Rencana Anggaran Biaya (RAB)",
        "    5.2 Jadwal Pelaksanaan (Time Schedule)",
        "DAFTAR PUSTAKA",
    ]
    for entry in toc_entries:
        add_para(entry, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
                 line_spacing=1.15, after=2)
    doc.add_page_break()

    # ----- ABSTRAK -----------------------------------------------------------
    add_para("ABSTRAK", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    for para in ABSTRAK_PARAGRAPHS:
        add_para(para, size=12, indent=1.0)
    add_para("")
    p = doc.add_paragraph()
    r1 = p.add_run("Kata Kunci: ")
    r1.bold = True
    r1.italic = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = p.add_run(KEYWORDS)
    r2.italic = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    doc.add_page_break()

    # ----- BODY SECTIONS -----------------------------------------------------
    for kind, heading, paras in SECTIONS:
        if kind == "BAB":
            # New page for each BAB
            add_para(heading, bold=True, size=14,
                     align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=12)
        elif kind == "H2":
            add_para(heading, bold=True, size=12,
                     align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=4)
            for para in paras:
                add_para(para, size=12, indent=1.0)
        elif kind == "H3":
            add_para(heading, bold=True, size=12,
                     align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=4)
        elif kind == "P":
            for para in paras:
                add_para(para, size=12, indent=1.0)
        elif kind == "OL":
            for idx, item in enumerate(paras, start=1):
                add_para(f"{idx}. {item}", size=12,
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line_spacing=1.5, after=4)
        elif kind == "UL":
            for item in paras:
                add_para(f"•  {item}", size=12,
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line_spacing=1.5, after=4)
        elif kind.startswith("TABLE_"):
            if heading:
                add_para(heading, bold=True, size=11,
                         align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4)
            data = TABLE_REGISTRY[kind]
            font_size = 9 if kind == "TABLE_SCHEDULE" else 11
            add_bordered_table(data, font_size=font_size)
            add_para("", after=6)

    # ----- DAFTAR PUSTAKA ----------------------------------------------------
    doc.add_page_break()
    add_para("DAFTAR PUSTAKA", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        pf = p.paragraph_format
        pf.line_spacing = 1.15
        pf.space_after = Pt(6)
        pf.left_indent = Cm(1)
        pf.first_line_indent = Cm(-1)  # hanging indent
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(out_path)
    print(f"DOCX written: {out_path}")


# =============================================================================
# PDF BUILDER (ReportLab)
# =============================================================================

def build_pdf(out_path: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
        KeepTogether,
    )

    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        leftMargin=4 * cm, rightMargin=3 * cm,
        topMargin=3 * cm, bottomMargin=3 * cm,
        title="Proposal Hilirisasi FSM LAPOR",
    )

    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "title", parent=base["Title"], fontName="Times-Bold",
            fontSize=14, leading=20, alignment=TA_CENTER, spaceAfter=8),
        "cover_title": ParagraphStyle(
            "cover_title", parent=base["Title"], fontName="Times-Bold",
            fontSize=14, leading=22, alignment=TA_CENTER, spaceAfter=12),
        "h1": ParagraphStyle(
            "h1", parent=base["Heading1"], fontName="Times-Bold",
            fontSize=14, leading=18, alignment=TA_CENTER, spaceAfter=10,
            spaceBefore=0),
        "h2": ParagraphStyle(
            "h2", parent=base["Heading2"], fontName="Times-Bold",
            fontSize=12, leading=16, spaceBefore=6, spaceAfter=4),
        "h3": ParagraphStyle(
            "h3", parent=base["Heading3"], fontName="Times-Bold",
            fontSize=12, leading=15, spaceBefore=4, spaceAfter=2),
        "body": ParagraphStyle(
            "body", parent=base["BodyText"], fontName="Times-Roman",
            fontSize=12, leading=18, alignment=TA_JUSTIFY,
            firstLineIndent=1 * cm, spaceAfter=4),
        "body_noindent": ParagraphStyle(
            "body_noindent", parent=base["BodyText"], fontName="Times-Roman",
            fontSize=12, leading=18, alignment=TA_JUSTIFY, spaceAfter=4),
        "list": ParagraphStyle(
            "list", parent=base["BodyText"], fontName="Times-Roman",
            fontSize=12, leading=18, alignment=TA_JUSTIFY,
            leftIndent=1 * cm, spaceAfter=4),
        "center": ParagraphStyle(
            "center", parent=base["BodyText"], fontName="Times-Roman",
            fontSize=12, leading=16, alignment=TA_CENTER, spaceAfter=4),
        "center_bold": ParagraphStyle(
            "center_bold", parent=base["BodyText"], fontName="Times-Bold",
            fontSize=12, leading=16, alignment=TA_CENTER, spaceAfter=4),
        "right": ParagraphStyle(
            "right", parent=base["BodyText"], fontName="Times-Roman",
            fontSize=12, leading=16, alignment=TA_RIGHT, spaceAfter=4),
        "ref": ParagraphStyle(
            "ref", parent=base["BodyText"], fontName="Times-Roman",
            fontSize=11, leading=14, alignment=TA_JUSTIFY,
            leftIndent=1 * cm, firstLineIndent=-1 * cm, spaceAfter=4),
        "table_caption": ParagraphStyle(
            "table_caption", parent=base["BodyText"], fontName="Times-Bold",
            fontSize=11, leading=14, alignment=TA_CENTER, spaceAfter=4,
            spaceBefore=6),
        "table_cell": ParagraphStyle(
            "table_cell", parent=base["BodyText"], fontName="Times-Roman",
            fontSize=10, leading=12, alignment=TA_LEFT),
        "table_cell_center": ParagraphStyle(
            "table_cell_center", parent=base["BodyText"],
            fontName="Times-Roman", fontSize=10, leading=12,
            alignment=TA_CENTER),
        "table_header": ParagraphStyle(
            "table_header", parent=base["BodyText"], fontName="Times-Bold",
            fontSize=10, leading=12, alignment=TA_CENTER),
        "table_cell_small": ParagraphStyle(
            "table_cell_small", parent=base["BodyText"],
            fontName="Times-Roman", fontSize=8, leading=10,
            alignment=TA_LEFT),
        "table_header_small": ParagraphStyle(
            "table_header_small", parent=base["BodyText"],
            fontName="Times-Bold", fontSize=8, leading=10,
            alignment=TA_CENTER),
    }

    flow = []

    def cell(text, style_name="table_cell"):
        return Paragraph(str(text).replace("&", "&amp;").replace(
            "<", "&lt;").replace(">", "&gt;"), styles[style_name])

    def make_table(rows, header=True, col_widths=None,
                   small_font=False, center_first_col=False):
        cell_style = "table_cell_small" if small_font else "table_cell"
        header_style = ("table_header_small" if small_font
                        else "table_header")
        center_style = "table_cell_center"
        data = []
        for i, row in enumerate(rows):
            new_row = []
            for j, val in enumerate(row):
                if header and i == 0:
                    new_row.append(cell(val, header_style))
                elif center_first_col and j == 0:
                    new_row.append(cell(val, center_style))
                else:
                    new_row.append(cell(val, cell_style))
            data.append(new_row)
        tbl = Table(data, colWidths=col_widths, repeatRows=1 if header else 0)
        ts = TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ])
        if header:
            ts.add("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9D9D9"))
        return Table(data, colWidths=col_widths,
                     repeatRows=1 if header else 0, style=ts)

    # ---- COVER --------------------------------------------------------------
    flow.append(Paragraph("PROPOSAL", styles["cover_title"]))
    flow.append(Paragraph("RISET PENUGASAN HILIRISASI",
                          styles["cover_title"]))
    flow.append(Spacer(1, 1.5 * cm))
    flow.append(Paragraph(TITLE, styles["cover_title"]))
    flow.append(Spacer(1, 1.5 * cm))
    flow.append(Paragraph("<b>Pengusul:</b>", styles["center"]))
    flow.append(Paragraph(PENGUSUL_NAME, styles["center"]))
    flow.append(Paragraph(PENGUSUL_NIP, styles["center"]))
    flow.append(Spacer(1, 4 * cm))
    flow.append(Paragraph("<b>FAKULTAS SAINS DAN MATEMATIKA</b>",
                          styles["center"]))
    flow.append(Paragraph("<b>UNIVERSITAS DIPONEGORO</b>",
                          styles["center"]))
    flow.append(Paragraph("<b>TAHUN 2026</b>", styles["center"]))
    flow.append(PageBreak())

    # ---- HALAMAN PENGESAHAN -------------------------------------------------
    flow.append(Paragraph("HALAMAN PENGESAHAN", styles["h1"]))
    pengesahan_data = []
    for no, label, value in PENGESAHAN_ROWS:
        pengesahan_data.append([
            cell(no, "table_cell"),
            cell(label, "table_cell"),
            cell(value, "table_cell"),
        ])
    pen_tbl = Table(
        pengesahan_data,
        colWidths=[1 * cm, 5.5 * cm, 7.5 * cm],
        style=TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]),
    )
    flow.append(pen_tbl)
    flow.append(Spacer(1, 0.8 * cm))
    flow.append(Paragraph("Semarang, ........... 2026", styles["right"]))
    flow.append(Spacer(1, 0.5 * cm))

    sig_data = [
        [Paragraph("Menyetujui,", styles["center"]),
         Paragraph("Ketua Pengusul,", styles["center"])],
        [Paragraph("Dekan FSM UNDIP", styles["center"]),
         Paragraph("", styles["center"])],
        [Paragraph("<br/><br/><br/>", styles["center"]),
         Paragraph("<br/><br/><br/>", styles["center"])],
        [Paragraph("<b>Prof. Dr. Kusworo Adi, S.Si., M.T.</b>",
                   styles["center"]),
         Paragraph("<b>Satriawan Rasyid Purnama, S.Kom., M.Cs.</b>",
                   styles["center"])],
        [Paragraph("NIP. 197203171998021001", styles["center"]),
         Paragraph("NIP. 199805212024061001", styles["center"])],
    ]
    sig_tbl = Table(sig_data, colWidths=[7 * cm, 7 * cm],
                    style=TableStyle([
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]))
    flow.append(sig_tbl)
    flow.append(PageBreak())

    # ---- DAFTAR ISI ---------------------------------------------------------
    flow.append(Paragraph("DAFTAR ISI", styles["h1"]))
    toc_entries = [
        "ABSTRAK",
        "BAB I PENDAHULUAN",
        "&nbsp;&nbsp;&nbsp;&nbsp;1.1 Latar Belakang",
        "&nbsp;&nbsp;&nbsp;&nbsp;1.2 Hasil Riset Awal (Sumber Hilirisasi)",
        "&nbsp;&nbsp;&nbsp;&nbsp;1.3 Tujuan Penelitian",
        "&nbsp;&nbsp;&nbsp;&nbsp;1.4 Manfaat Penelitian",
        "BAB II MITRA DAN TARGET PENGGUNAAN",
        "&nbsp;&nbsp;&nbsp;&nbsp;2.1 Mitra untuk Hilirisasi",
        "&nbsp;&nbsp;&nbsp;&nbsp;2.2 Target Penggunaan",
        "&nbsp;&nbsp;&nbsp;&nbsp;2.3 Dampak dan Manfaat Hilirisasi",
        "BAB III METODOLOGI HILIRISASI",
        "&nbsp;&nbsp;&nbsp;&nbsp;3.1 Tahapan Kegiatan",
        "&nbsp;&nbsp;&nbsp;&nbsp;3.2 Metode Pengujian dan Validasi",
        "&nbsp;&nbsp;&nbsp;&nbsp;3.3 Indikator Kinerja dan Ketercapaian",
        "BAB IV RENCANA LUARAN",
        "&nbsp;&nbsp;&nbsp;&nbsp;4.1 Luaran",
        "BAB V RENCANA ANGGARAN DAN JADWAL PELAKSANAAN",
        "&nbsp;&nbsp;&nbsp;&nbsp;5.1 Rencana Anggaran Biaya (RAB)",
        "&nbsp;&nbsp;&nbsp;&nbsp;5.2 Jadwal Pelaksanaan (Time Schedule)",
        "DAFTAR PUSTAKA",
    ]
    for entry in toc_entries:
        flow.append(Paragraph(entry, styles["body_noindent"]))
    flow.append(PageBreak())

    # ---- ABSTRAK ------------------------------------------------------------
    flow.append(Paragraph("ABSTRAK", styles["h1"]))
    for para in ABSTRAK_PARAGRAPHS:
        flow.append(Paragraph(para, styles["body"]))
    flow.append(Spacer(1, 6))
    flow.append(Paragraph(
        f"<b><i>Kata Kunci:</i></b> <i>{KEYWORDS}</i>",
        styles["body_noindent"]))
    flow.append(PageBreak())

    # ---- BODY ---------------------------------------------------------------
    content_width = A4[0] - 4 * cm - 3 * cm  # ~14 cm
    rab_widths = [1 * cm, content_width - 1 * cm - 1.5 * cm - 1.5 * cm
                  - 2.5 * cm - 2.5 * cm,
                  1.5 * cm, 1.5 * cm, 2.5 * cm, 2.5 * cm]
    role_widths = [2 * cm, 4 * cm, content_width - 6 * cm]
    kpi_widths = [0.8 * cm, 5 * cm, content_width - 5.8 * cm]
    produk_widths = [0.8 * cm, 5 * cm, content_width - 5.8 * cm]
    luaran_widths = [3 * cm, content_width - 3 * cm]
    rekap_widths = [0.8 * cm, 6 * cm, 3.5 * cm, content_width - 10.3 * cm]
    sched_no = 0.8 * cm
    sched_kegiatan = 5 * cm
    sched_month = (content_width - sched_no - sched_kegiatan) / 12
    sched_widths = [sched_no, sched_kegiatan] + [sched_month] * 12

    table_widths = {
        "TABLE_ROLE": role_widths,
        "TABLE_KPI": kpi_widths,
        "TABLE_PRODUK": produk_widths,
        "TABLE_LUARAN": luaran_widths,
        "TABLE_RAB_A": rab_widths,
        "TABLE_RAB_B": rab_widths,
        "TABLE_RAB_REKAP": rekap_widths,
        "TABLE_SCHEDULE": sched_widths,
    }

    for kind, heading, paras in SECTIONS:
        if kind == "BAB":
            flow.append(PageBreak())
            flow.append(Paragraph(heading, styles["h1"]))
        elif kind == "H2":
            flow.append(Paragraph(heading, styles["h2"]))
            for para in paras:
                flow.append(Paragraph(para, styles["body"]))
        elif kind == "H3":
            flow.append(Paragraph(heading, styles["h3"]))
        elif kind == "P":
            for para in paras:
                flow.append(Paragraph(para, styles["body"]))
        elif kind == "OL":
            for idx, item in enumerate(paras, start=1):
                flow.append(Paragraph(f"{idx}. {item}", styles["list"]))
        elif kind == "UL":
            for item in paras:
                flow.append(Paragraph(f"•&nbsp;&nbsp;{item}",
                                      styles["list"]))
        elif kind.startswith("TABLE_"):
            if heading:
                flow.append(Paragraph(heading, styles["table_caption"]))
            data = TABLE_REGISTRY[kind]
            small = (kind == "TABLE_SCHEDULE")
            tbl = make_table(
                data, header=True, col_widths=table_widths.get(kind),
                small_font=small,
                center_first_col=(kind in ("TABLE_KPI", "TABLE_PRODUK",
                                            "TABLE_RAB_A", "TABLE_RAB_B",
                                            "TABLE_RAB_REKAP",
                                            "TABLE_SCHEDULE")),
            )
            flow.append(tbl)
            flow.append(Spacer(1, 8))

    # ---- DAFTAR PUSTAKA -----------------------------------------------------
    flow.append(PageBreak())
    flow.append(Paragraph("DAFTAR PUSTAKA", styles["h1"]))
    for ref in REFERENCES:
        flow.append(Paragraph(ref, styles["ref"]))

    doc.build(flow)
    print(f"PDF written: {out_path}")


# =============================================================================
# MAIN
# =============================================================================

def main() -> None:
    here = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(here, "proposal-hilirisasi-fsm-lapor.docx")
    pdf_path = os.path.join(here, "proposal-hilirisasi-fsm-lapor.pdf")

    build_docx(docx_path)
    build_pdf(pdf_path)


if __name__ == "__main__":
    main()
