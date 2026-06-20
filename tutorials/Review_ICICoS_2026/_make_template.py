"""
_make_template.py
Generates review-template.docx — representative ICICoS 2026 review form.
Run once: python3 _make_template.py
NOTE: Replace this file with the REAL template when provided by the committee.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name="Times New Roman", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    size = 14 if level == 1 else 11
    set_font(run, bold=True, size=size)
    return p


def add_divider(doc):
    p = doc.add_paragraph("─" * 68)
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    p.runs[0].font.size = Pt(8)
    return p


def add_label_field(doc, label, placeholder="[...]", inline=True):
    """Add a labeled field. inline=True → label: [...] on same line."""
    p = doc.add_paragraph()
    r_label = p.add_run(label)
    set_font(r_label, bold=True, size=11)
    if inline:
        r_sep = p.add_run("  ")
        set_font(r_sep, size=11)
        r_ph = p.add_run(placeholder)
        set_font(r_ph, size=11)
    return p


def add_block_field(doc, label, placeholder="[...]"):
    """Add label paragraph then placeholder paragraph below it."""
    p_label = doc.add_paragraph()
    r = p_label.add_run(label)
    set_font(r, bold=True, size=11)

    p_content = doc.add_paragraph()
    r2 = p_content.add_run(placeholder)
    set_font(r2, size=11)
    p_content.paragraph_format.left_indent = Inches(0.25)
    return p_label, p_content


def add_decision_field(doc, label, options, separator=" / "):
    """Add a numbered label then the options on the next line."""
    p_label = doc.add_paragraph()
    r = p_label.add_run(label)
    set_font(r, bold=True, size=11)

    p_opts = doc.add_paragraph()
    p_opts.paragraph_format.left_indent = Inches(0.25)
    for i, opt in enumerate(options):
        r_opt = p_opts.add_run(opt)
        set_font(r_opt, size=11, bold=False)
        if i < len(options) - 1:
            r_sep = p_opts.add_run(separator)
            set_font(r_sep, size=11)
    return p_label, p_opts


# ─────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── HEADER ──────────────────────────────────
add_heading(doc, "REVIEW FORM — ICICoS 2026")
add_heading(doc, "International Conference on Information and Communication Technology", level=2)
add_divider(doc)

# ── PAPER IDENTITY ──────────────────────────
p = doc.add_paragraph()
r = p.add_run("PAPER IDENTITY")
set_font(r, bold=True, size=10, color=(0x55, 0x55, 0x55))

add_label_field(doc, "Paper ID      :", inline=True)
add_label_field(doc, "Paper Title   :", inline=True)
add_label_field(doc, "Reviewer ID   :", inline=True)
add_divider(doc)

# ── SCORES ──────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("EVALUATION SCORES   (1 = Very Poor  ·  3 = Acceptable  ·  5 = Excellent)")
set_font(r, bold=True, size=10, color=(0x55, 0x55, 0x55))

add_label_field(doc, "1.  Originality / Novelty          :", inline=True)
add_label_field(doc, "2.  Technical Quality              :", inline=True)
add_label_field(doc, "3.  Clarity of Presentation        :", inline=True)
add_label_field(doc, "4.  Relevance to ICICoS            :", inline=True)
add_label_field(doc, "5.  Quality of References          :", inline=True)
add_label_field(doc, "6.  Overall Score                  :", inline=True)
add_divider(doc)

# ── TEXT FIELDS ─────────────────────────────
p = doc.add_paragraph()
r = p.add_run("REVIEW COMMENTS")
set_font(r, bold=True, size=10, color=(0x55, 0x55, 0x55))

add_block_field(doc, "7.  Summary of the Paper")
doc.add_paragraph()   # spacer

add_block_field(doc, "8.  Strengths")
doc.add_paragraph()

add_block_field(doc, "9.  Weaknesses and Suggestions for Improvement")
doc.add_paragraph()

add_block_field(doc, "10. Questions for Authors  (optional — leave blank if none)")
doc.add_paragraph()

add_block_field(doc, "11. Confidential Comments to Editor  (optional — not shown to authors)")
add_divider(doc)

# ── DECISION ────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("DECISION")
set_font(r, bold=True, size=10, color=(0x55, 0x55, 0x55))

add_decision_field(doc, "12. Overall Recommendation",
                   ["Accept", "Minor Revision", "Major Revision", "Reject"])
doc.add_paragraph()

add_decision_field(doc, "13. Reviewer Confidence",
                   ["Expert", "Knowledgeable", "Passing Knowledge", "Basic"])
add_divider(doc)

# ── FOOTER ──────────────────────────────────
p = doc.add_paragraph("Thank you for your contribution to ICICoS 2026.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.runs[0], size=9, color=(0x88, 0x88, 0x88))

# ── SAVE ────────────────────────────────────
doc.save("review-template.docx")
print("✅ review-template.docx created.")
