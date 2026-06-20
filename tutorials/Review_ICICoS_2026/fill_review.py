"""
fill_review.py — ICICoS 2026 Review Template Filler
=====================================================
Mengisi review-template.docx dengan konten review:
  • Mengganti setiap [...]  dengan teks/nilai yang sesuai
  • Menebalkan (bold) pilihan yang dipilih pada item keputusan (item 12 & 13)

Dependensi : pip install python-docx
Penggunaan : python3 fill_review.py
Output     : review-filled.docx  (template asli tidak berubah)
"""

from docx import Document


# ══════════════════════════════════════════════════════
#  BAGIAN 1 — ISI DATA REVIEW DI SINI
# ══════════════════════════════════════════════════════

review_data = {
    # ── Identitas ─────────────────────────────────────
    "paper_id"        : "ICICoS-2026-042",
    "paper_title"     : "Adaptive Deep Learning for Earthquake Early Warning",
    "reviewer_id"     : "R-07",

    # ── Skor (angka 1–5 sebagai string) ───────────────
    "score_originality" : "4",
    "score_technical"   : "3",
    "score_clarity"     : "4",
    "score_relevance"   : "5",
    "score_references"  : "3",
    "score_overall"     : "4",

    # ── Komentar (bahasa Inggris formal) ──────────────
    "summary"     : (
        "This paper proposes an adaptive deep learning framework for seismic "
        "early warning that dynamically adjusts model architecture based on "
        "incoming signal quality. The authors evaluate their approach on two "
        "public datasets and report improvements over conventional CNN baselines."
    ),
    "strengths"   : (
        "1. The adaptive signal-quality gating mechanism is a genuine novelty "
        "not observed in prior ICICoS submissions.\n"
        "2. Results on both STEAD and INSTANCE datasets are consistent and "
        "clearly presented in Table 2.\n"
        "3. The writing is well-structured and the introduction provides "
        "strong motivation."
    ),
    "weaknesses"  : (
        "1. Baseline comparison (Section 4.2) omits PhaseNet and EQTransformer, "
        "which are standard references in this domain. Please add or justify "
        "their exclusion.\n"
        "2. No ablation study is provided; the contribution of the gating "
        "mechanism versus the backbone cannot be isolated. A table reporting "
        "ablated variants is strongly recommended.\n"
        "3. All experiments use a single random seed. Please report mean ± std "
        "over at least 3 runs."
    ),
    "questions"   : (
        "1. How does the gating mechanism behave when signal quality degrades "
        "gradually versus abruptly? Is there a latency penalty?\n"
        "2. What is the inference time on a CPU-only edge device, given the "
        "real-time constraint of EEW systems?"
    ),
    "confidential": "",   # kosongkan jika tidak ada

    # ── Keputusan ─────────────────────────────────────
    # Pilih PERSIS salah satu opsi yang ada di template:
    # "Accept"  |  "Minor Revision"  |  "Major Revision"  |  "Reject"
    "decision"    : "Minor Revision",

    # "Expert"  |  "Knowledgeable"  |  "Passing Knowledge"  |  "Basic"
    "confidence"  : "Knowledgeable",
}


# ══════════════════════════════════════════════════════
#  BAGIAN 2 — KONFIGURASI (sesuaikan jika template ganti)
# ══════════════════════════════════════════════════════

TEMPLATE_PATH = "review-template.docx"
OUTPUT_PATH   = "review-filled.docx"

# Kata kunci untuk mendeteksi field dari teks paragraf label.
# Format: "substring_unik_di_label" → "key_di_review_data"
INLINE_FIELDS = {
    # label dan [...] ada dalam satu paragraf
    "Paper ID"          : "paper_id",
    "Paper Title"       : "paper_title",
    "Reviewer ID"       : "reviewer_id",
    "Originality"       : "score_originality",
    "Technical Quality" : "score_technical",
    "Clarity"           : "score_clarity",
    "Relevance"         : "score_relevance",
    "Quality of Ref"    : "score_references",
    "Overall Score"     : "score_overall",
}

BLOCK_LABELS = {
    # label pada paragraf N, [...] pada paragraf N+1
    "Summary of the Paper"       : "summary",
    "Strengths"                  : "strengths",
    "Weaknesses"                 : "weaknesses",
    "Questions for Authors"      : "questions",
    "Confidential Comments"      : "confidential",
}

DECISION_LABELS = {
    # opsi Accept/Reject/dll pada paragraf setelah label ini
    "Overall Recommendation" : "decision",
    "Reviewer Confidence"    : "confidence",
}


# ══════════════════════════════════════════════════════
#  BAGIAN 3 — FUNGSI INTI
# ══════════════════════════════════════════════════════

def fill_inline(para, value):
    """
    Ganti [...] dalam paragraf yang label dan placeholder-nya satu baris.
    Contoh: 'Paper ID :  [...]'  →  'Paper ID :  ICICoS-2026-042'
    Hanya mengubah run yang berisi '[...]', run lain (label bold) tidak disentuh.
    """
    for run in para.runs:
        if "[...]" in run.text:
            run.text = run.text.replace("[...]", value)
            return True
    return False


def fill_block(para, value):
    """
    Ganti seluruh teks paragraf standalone '[...]' dengan nilai.
    Mempertahankan formatting run pertama.
    """
    full = "".join(r.text for r in para.runs)
    if "[...]" not in full:
        return False
    # Tulis di run pertama, kosongkan sisanya
    if para.runs:
        para.runs[0].text = value
        for r in para.runs[1:]:
            r.text = ""
    return True


def bold_selection(para, selected):
    """
    Bold satu opsi dalam baris keputusan.
    Contoh: 'Accept / Minor Revision / Major Revision / Reject'
    → run 'Minor Revision' menjadi bold, sisanya bold=False.
    Fungsi ini hanya mengubah atribut .bold; teks tidak berubah.
    """
    changed = False
    for run in para.runs:
        text = run.text.strip()
        if not text or text == "/":
            continue
        run.bold = (text == selected)
        changed = True
    return changed


# ══════════════════════════════════════════════════════
#  BAGIAN 4 — PROSES DOKUMEN
# ══════════════════════════════════════════════════════

def process(template_path, output_path, data):
    doc = Document(template_path)
    paragraphs = doc.paragraphs
    n = len(paragraphs)

    pending_block   = None   # key untuk block field yang menunggu [...]
    pending_decision = None  # key untuk decision line yang akan datang

    for i, para in enumerate(paragraphs):
        text = para.text.strip()

        # ── 1. Cek apakah ini label untuk BLOCK field ──────────
        if pending_block:
            if "[...]" in para.text:
                value = data.get(pending_block, "")
                fill_block(para, value)
            pending_block = None
            continue

        # ── 2. Cek apakah ini baris OPTIONS untuk DECISION ─────
        if pending_decision:
            selected = data.get(pending_decision, "")
            bold_selection(para, selected)
            pending_decision = None
            continue

        # ── 3. Cek INLINE fields (label + [...] satu baris) ────
        matched_inline = False
        for keyword, field_key in INLINE_FIELDS.items():
            if keyword in text and "[...]" in text:
                fill_inline(para, data.get(field_key, ""))
                matched_inline = True
                break
        if matched_inline:
            continue

        # ── 4. Cek BLOCK label (set pending untuk paragraf berikutnya) ──
        for keyword, field_key in BLOCK_LABELS.items():
            if keyword in text:
                pending_block = field_key
                break

        # ── 5. Cek DECISION label (set pending untuk baris opsi) ────────
        for keyword, field_key in DECISION_LABELS.items():
            if keyword in text:
                pending_decision = field_key
                break

    doc.save(output_path)

    # ── Laporan ──────────────────────────────────────────────────
    print(f"\n✅  Tersimpan: {output_path}")
    print(f"    Decision   : {data['decision']}  (ditebalkan di item 12)")
    print(f"    Confidence : {data['confidence']}  (ditebalkan di item 13)")
    print(f"    Summary    : {data['summary'][:60]}…")


# ══════════════════════════════════════════════════════
#  BAGIAN 5 — VERIFIKASI OTOMATIS (self-check)
# ══════════════════════════════════════════════════════

def verify(output_path, data):
    """
    Baca output dan pastikan tidak ada [...] yang tertinggal,
    serta decision sudah ditebalkan dengan benar.
    """
    doc = Document(output_path)
    errors = []
    decision_ok = False
    confidence_ok = False

    for para in doc.paragraphs:
        text = para.text

        # Cek sisa [...]
        if "[...]" in text:
            errors.append(f"  ⚠  Placeholder belum terisi: '{text[:60]}'")

        # Cek bold decision
        for run in para.runs:
            if run.text.strip() == data["decision"] and run.bold:
                decision_ok = True
            if run.text.strip() == data["confidence"] and run.bold:
                confidence_ok = True

    print("\n── Verifikasi ──────────────────────────────────────────")
    if errors:
        for e in errors:
            print(e)
    else:
        print("  ✓  Tidak ada [...] yang tertinggal")

    if decision_ok:
        print(f"  ✓  Decision '{data['decision']}' ditebalkan")
    else:
        print(f"  ✗  Decision '{data['decision']}' TIDAK ditebalkan — cek DECISION_LABELS")

    if confidence_ok:
        print(f"  ✓  Confidence '{data['confidence']}' ditebalkan")
    else:
        print(f"  ✗  Confidence '{data['confidence']}' TIDAK ditebalkan — cek DECISION_LABELS")


# ══════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════

if __name__ == "__main__":
    process(TEMPLATE_PATH, OUTPUT_PATH, review_data)
    verify(OUTPUT_PATH, review_data)
